"""
resume/builder.py — Resume Builder
═══════════════════════════════════════════════════════════════
Generates DOCX and PDF resumes from structured ResumeData.
ATS-friendly formatting, professional layout.

Interface:
  ResumeBuilder()
  ResumeBuilder.build_docx(resume_data, template_path?, company?, title?) → str
  ResumeBuilder.build_pdf(resume_data, company?, title?) → str
  ResumeBuilder.build_both(resume_data, template_path?, company?, title?) → dict
  ResumeBuilder.get_output_path(company, title) → str

Dependencies: config.py, core/logger.py, profile/resume_data.py
Libraries: python-docx, weasyprint (optional for PDF)
"""

import os
import sys
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, List, Any, Union

# ─── Project root on path ──────────────────────────────────────
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import RESUME_OUTPUT_DIR, RESUME_CONFIG, BASE_DIR
from core.logger import get_logger

logger = get_logger("resume.builder")


# ════════════════════════════════════════════════════════════════
#  RESUME BUILDER
# ════════════════════════════════════════════════════════════════

class ResumeBuilder:
    """
    Builds ATS-friendly DOCX and PDF resumes from structured ResumeData.
    
    - DOCX via python-docx (primary output)
    - PDF via weasyprint (HTML → PDF, optional fallback)
    - Template support (fills placeholders in existing DOCX)
    - Clean professional formatting with section headers, bullets
    - Output organized: resume/output/COMPANY_TITLE_DATE/
    """

    # ── Color palette (consistent branding) ─────────────────────
    COLOR_PRIMARY = (0x1A, 0x1A, 0x2E)      # Dark navy — headers
    COLOR_SECONDARY = (0x44, 0x44, 0x44)     # Dark gray — company names
    COLOR_BODY = (0x33, 0x33, 0x33)          # Body text
    COLOR_META = (0x77, 0x77, 0x77)          # Dates, metadata
    COLOR_CONTACT = (0x55, 0x55, 0x55)       # Contact info
    COLOR_ACCENT = (0x2E, 0x7D, 0x32)        # Green — recognition/stars
    COLOR_LINK = (0x29, 0x62, 0xFF)          # Blue — links
    COLOR_RULE = (0xCC, 0xCC, 0xCC)          # Light gray — dividers
    COLOR_PRIMARY_HEX = '1A1A2E'
    COLOR_RULE_HEX = 'CCCCCC'

    def __init__(self):
        """Initialize builder, create output directory."""
        self.output_dir = Path(RESUME_OUTPUT_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.config = RESUME_CONFIG
        logger.info(f"ResumeBuilder initialized — output: {self.output_dir}")

    # ════════════════════════════════════════════════════════════
    #  PUBLIC API
    # ════════════════════════════════════════════════════════════

    def get_output_path(self, company: str, title: str) -> str:
        """
        Generate organized output directory path.
        
        Format: resume/output/COMPANY_TITLE_YYYYMMDD_HHMMSS/
        
        Args:
            company: Target company name
            title: Target job title
            
        Returns:
            str: Absolute path to output directory (created)
        """
        company_clean = self._sanitize_filename(company)
        title_clean = self._sanitize_filename(title)
        date_str = datetime.now().strftime("%Y%m%d_%H%M%S")

        folder_name = f"{company_clean}_{title_clean}_{date_str}"
        output_path = self.output_dir / folder_name
        output_path.mkdir(parents=True, exist_ok=True)

        return str(output_path)

    def build_docx(
        self,
        resume_data,
        template_path: Optional[str] = None,
        company: Optional[str] = None,
        title: Optional[str] = None,
    ) -> str:
        """
        Build a DOCX resume from ResumeData.
        
        Args:
            resume_data: ResumeData dataclass instance or dict
            template_path: Optional path to base .docx template with placeholders
            company: Target company name (for file naming)
            title: Target job title (for file naming)
            
        Returns:
            str: Absolute path to generated .docx file
        """
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise ImportError("python-docx required. Install: pip install python-docx")

        data = self._to_dict(resume_data)

        # Resolve naming
        company = company or data.get('_target_company', 'General')
        title = title or data.get('_target_title', 'Resume')

        output_path = self.get_output_path(company, title)
        docx_path = os.path.join(output_path, "resume.docx")

        # Build document
        if template_path and os.path.exists(template_path):
            logger.info(f"Building DOCX from template: {template_path}")
            doc = self._build_from_template(data, template_path)
        else:
            logger.info("Building fresh DOCX resume")
            doc = self._build_fresh_docx(data)

        doc.save(docx_path)
        logger.info(f"✅ DOCX saved: {docx_path}")
        return docx_path

    def build_pdf(
        self,
        resume_data,
        company: Optional[str] = None,
        title: Optional[str] = None,
    ) -> str:
        """
        Build a PDF resume from ResumeData via HTML→weasyprint.
        
        Falls back to saving HTML if weasyprint not installed.
        
        Args:
            resume_data: ResumeData dataclass instance or dict
            company: Target company name
            title: Target job title
            
        Returns:
            str: Path to generated .pdf (or .html fallback)
        """
        data = self._to_dict(resume_data)

        company = company or data.get('_target_company', 'General')
        title = title or data.get('_target_title', 'Resume')

        output_path = self.get_output_path(company, title)
        pdf_path = os.path.join(output_path, "resume.pdf")
        html_path = os.path.join(output_path, "resume.html")

        # Generate HTML
        html_content = self._generate_html(data)

        # Always save HTML (useful for debug and as fallback)
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        logger.info(f"HTML saved: {html_path}")

        # Convert to PDF
        try:
            from weasyprint import HTML as WeasyHTML
            WeasyHTML(string=html_content).write_pdf(pdf_path)
            logger.info(f"✅ PDF saved: {pdf_path}")
            return pdf_path
        except ImportError:
            logger.warning(
                "weasyprint not installed — PDF skipped. "
                "Install: pip install weasyprint (needs system deps)"
            )
            return html_path
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            return html_path

    def build_both(
        self,
        resume_data,
        template_path: Optional[str] = None,
        company: Optional[str] = None,
        title: Optional[str] = None,
    ) -> Dict[str, Optional[str]]:
        """
        Build both DOCX and PDF versions.
        
        Returns:
            dict: {'docx': path_or_None, 'pdf': path_or_None}
        """
        result = {'docx': None, 'pdf': None}

        try:
            result['docx'] = self.build_docx(resume_data, template_path, company, title)
        except Exception as e:
            logger.error(f"DOCX build failed: {e}", exc_info=True)

        try:
            result['pdf'] = self.build_pdf(resume_data, company, title)
        except Exception as e:
            logger.error(f"PDF build failed: {e}", exc_info=True)

        return result

    # ════════════════════════════════════════════════════════════
    #  DOCX BUILDER — FRESH (from scratch)
    # ════════════════════════════════════════════════════════════

    def _build_fresh_docx(self, data: dict):
        """Build a clean, ATS-optimized DOCX resume from scratch."""
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # ── Page margins (tight but readable) ──
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)

        # ── Setup default styles ──
        self._setup_styles(doc)

        # ── HEADER: Name ──
        name = data.get('name', 'Name')
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.space_before = Pt(0)
        name_para.space_after = Pt(2)
        name_run = name_para.add_run(name.upper())
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_run.font.color.rgb = RGBColor(*self.COLOR_PRIMARY)

        # ── Contact line ──
        contact_parts = []
        if data.get('email'):
            contact_parts.append(data['email'])
        if data.get('phone'):
            contact_parts.append(data['phone'])
        if data.get('location'):
            contact_parts.append(data['location'])

        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.space_before = Pt(0)
            contact_para.space_after = Pt(1)
            c_run = contact_para.add_run(' │ '.join(contact_parts))
            c_run.font.size = Pt(9)
            c_run.font.color.rgb = RGBColor(*self.COLOR_CONTACT)

        # ── Links line ──
        link_parts = []
        if data.get('linkedin'):
            link_parts.append(f"LinkedIn: {data['linkedin']}")
        if data.get('github'):
            link_parts.append(f"GitHub: {data['github']}")

        if link_parts:
            link_para = doc.add_paragraph()
            link_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            link_para.space_before = Pt(0)
            link_para.space_after = Pt(4)
            l_run = link_para.add_run(' │ '.join(link_parts))
            l_run.font.size = Pt(9)
            l_run.font.color.rgb = RGBColor(*self.COLOR_CONTACT)

        # ── Divider ──
        self._add_thin_rule(doc, self.COLOR_RULE_HEX)

        # ── PROFESSIONAL SUMMARY ──
        if data.get('summary'):
            self._add_section_header(doc, 'PROFESSIONAL SUMMARY')
            s_para = doc.add_paragraph()
            s_para.space_before = Pt(2)
            s_para.space_after = Pt(6)
            s_run = s_para.add_run(data['summary'])
            s_run.font.size = Pt(10)
            s_run.font.color.rgb = RGBColor(*self.COLOR_BODY)

        # ── TECHNICAL SKILLS ──
        skills = data.get('skills', {})
        if skills:
            self._add_section_header(doc, 'TECHNICAL SKILLS')
            self._add_skills_section(doc, skills)

        # ── WORK EXPERIENCE ──
        experience = data.get('experience', [])
        if experience:
            self._add_section_header(doc, 'WORK EXPERIENCE')
            for i, exp in enumerate(experience):
                self._add_experience_entry(doc, exp)
                if i < len(experience) - 1:
                    sp = doc.add_paragraph()
                    sp.space_before = Pt(0)
                    sp.space_after = Pt(2)

        # ── PROJECTS ──
        projects = data.get('projects', [])
        if projects:
            self._add_section_header(doc, 'PROJECTS')
            for i, proj in enumerate(projects):
                self._add_project_entry(doc, proj)
                if i < len(projects) - 1:
                    sp = doc.add_paragraph()
                    sp.space_before = Pt(0)
                    sp.space_after = Pt(2)

        # ── EDUCATION ──
        education = data.get('education', [])
        if education:
            self._add_section_header(doc, 'EDUCATION')
            for edu in education:
                self._add_education_entry(doc, edu)

        # ── CERTIFICATIONS ──
        certifications = data.get('certifications', [])
        if certifications:
            self._add_section_header(doc, 'CERTIFICATIONS')
            for cert in certifications:
                cert_text = self._cert_to_text(cert)
                if cert_text:
                    cp = doc.add_paragraph(style='List Bullet')
                    cp.space_after = Pt(1)
                    cp.space_before = Pt(0)
                    cr = cp.add_run(cert_text)
                    cr.font.size = Pt(10)

        # ── ACHIEVEMENTS ──
        achievements = data.get('achievements', [])
        if achievements:
            self._add_section_header(doc, 'ACHIEVEMENTS')
            for ach in achievements:
                ach_text = str(ach) if not isinstance(ach, str) else ach
                ap = doc.add_paragraph(style='List Bullet')
                ap.space_after = Pt(1)
                ap.space_before = Pt(0)
                ar = ap.add_run(ach_text)
                ar.font.size = Pt(10)

        # ── CODING PROFILES ──
        coding_profiles = data.get('coding_profiles', {})
        if coding_profiles:
            self._add_section_header(doc, 'CODING PROFILES')
            profile_parts = []
            for platform, detail in coding_profiles.items():
                if isinstance(detail, str):
                    profile_parts.append(f"{platform}: {detail}")
                elif isinstance(detail, dict):
                    desc = detail.get('description', detail.get('url', ''))
                    profile_parts.append(f"{platform}: {desc}")
            if profile_parts:
                pp = doc.add_paragraph()
                pp.space_after = Pt(2)
                pr = pp.add_run(' │ '.join(profile_parts))
                pr.font.size = Pt(10)

        return doc

    # ════════════════════════════════════════════════════════════
    #  DOCX BUILDER — FROM TEMPLATE
    # ════════════════════════════════════════════════════════════

    def _build_from_template(self, data: dict, template_path: str):
        """
        Build resume from existing DOCX template.
        
        Replaces placeholders like {{NAME}}, {{SUMMARY}}, etc.
        Preserves template formatting, fonts, layout.
        """
        from docx import Document

        doc = Document(template_path)

        # ── Define all placeholder mappings ──
        placeholders = {
            '{{NAME}}': data.get('name', ''),
            '{{EMAIL}}': data.get('email', ''),
            '{{PHONE}}': data.get('phone', ''),
            '{{LOCATION}}': data.get('location', ''),
            '{{LINKEDIN}}': data.get('linkedin', ''),
            '{{GITHUB}}': data.get('github', ''),
            '{{SUMMARY}}': data.get('summary', ''),
        }

        # Build skills text
        skills = data.get('skills', {})
        if isinstance(skills, dict):
            skills_lines = []
            for cat, items in skills.items():
                if items:
                    items_str = ', '.join(items) if isinstance(items, list) else str(items)
                    skills_lines.append(f"{cat}: {items_str}")
            placeholders['{{SKILLS}}'] = '\n'.join(skills_lines)
        elif isinstance(skills, list):
            placeholders['{{SKILLS}}'] = ', '.join(skills)
        else:
            placeholders['{{SKILLS}}'] = ''

        # Build experience text
        exp_lines = []
        for exp in data.get('experience', []):
            title = exp.get('title', '')
            company = exp.get('company', '')
            duration = exp.get('duration', '')
            exp_lines.append(f"{title} — {company} ({duration})")
            for bullet in self._extract_bullets(exp):
                exp_lines.append(f"  • {bullet}")
            exp_lines.append('')
        placeholders['{{EXPERIENCE}}'] = '\n'.join(exp_lines)

        # Build projects text
        proj_lines = []
        for proj in data.get('projects', []):
            name = proj.get('name', '')
            tech = proj.get('tech', proj.get('technologies', ''))
            if isinstance(tech, list):
                tech = ', '.join(tech)
            proj_lines.append(f"{name} ({tech})")
            for bullet in self._extract_bullets(proj):
                proj_lines.append(f"  • {bullet}")
            proj_lines.append('')
        placeholders['{{PROJECTS}}'] = '\n'.join(proj_lines)

        # Build education text
        edu_lines = []
        for edu in data.get('education', []):
            degree = edu.get('degree', '')
            uni = edu.get('university', edu.get('institution', ''))
            year = edu.get('year', '')
            cgpa = edu.get('cgpa', '')
            line = f"{degree} — {uni}"
            if year:
                line += f" ({year})"
            if cgpa:
                line += f" | CGPA: {cgpa}"
            edu_lines.append(line)
        placeholders['{{EDUCATION}}'] = '\n'.join(edu_lines)

        # Certifications text
        cert_lines = [self._cert_to_text(c) for c in data.get('certifications', []) if c]
        placeholders['{{CERTIFICATIONS}}'] = '\n'.join([f"• {c}" for c in cert_lines if c])

        # Achievements text
        ach_lines = [str(a) for a in data.get('achievements', []) if a]
        placeholders['{{ACHIEVEMENTS}}'] = '\n'.join([f"• {a}" for a in ach_lines])

        # ── Replace in all paragraphs ──
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, placeholders)

        # ── Replace in all tables ──
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, placeholders)

        # ── Replace in headers/footers ──
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    self._replace_in_paragraph(para, placeholders)
            if section.footer:
                for para in section.footer.paragraphs:
                    self._replace_in_paragraph(para, placeholders)

        logger.info(f"Template-based resume built from: {template_path}")
        return doc

    @staticmethod
    def _replace_in_paragraph(para, placeholders: dict):
        """Replace placeholders in a paragraph while preserving formatting."""
        full_text = para.text
        if not any(ph in full_text for ph in placeholders):
            return

        for placeholder, value in placeholders.items():
            if placeholder not in full_text:
                continue

            # Try run-level replacement first (preserves formatting)
            for run in para.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

            # If placeholder spanned multiple runs, do paragraph-level
            if placeholder in para.text:
                # Rebuild runs — fallback
                combined = para.text
                combined = combined.replace(placeholder, value)
                # Clear all runs and set first run
                if para.runs:
                    para.runs[0].text = combined
                    for run in para.runs[1:]:
                        run.text = ''

    # ════════════════════════════════════════════════════════════
    #  DOCX HELPER — SECTION COMPONENTS
    # ════════════════════════════════════════════════════════════

    def _setup_styles(self, doc):
        """Configure default document styles."""
        from docx.shared import Pt, RGBColor

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)
        font.color.rgb = RGBColor(*self.COLOR_BODY)

        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.15

    def _add_section_header(self, doc, title: str):
        """Add section header with bottom border line."""
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        para = doc.add_paragraph()
        para.space_before = Pt(10)
        para.space_after = Pt(4)

        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(*self.COLOR_PRIMARY)

        # Bottom border on paragraph
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), self.COLOR_PRIMARY_HEX)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_thin_rule(self, doc, color_hex: str = 'CCCCCC'):
        """Add a thin horizontal divider."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Pt

        para = doc.add_paragraph()
        para.space_before = Pt(0)
        para.space_after = Pt(0)

        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_skills_section(self, doc, skills):
        """Add technical skills section."""
        from docx.shared import Pt

        if isinstance(skills, dict):
            for category, skill_list in skills.items():
                if not skill_list:
                    continue
                sp = doc.add_paragraph()
                sp.space_before = Pt(1)
                sp.space_after = Pt(2)

                cat_run = sp.add_run(f"{category}: ")
                cat_run.bold = True
                cat_run.font.size = Pt(10)

                if isinstance(skill_list, list):
                    skills_text = ', '.join(str(s) for s in skill_list)
                else:
                    skills_text = str(skill_list)

                val_run = sp.add_run(skills_text)
                val_run.font.size = Pt(10)

        elif isinstance(skills, list):
            sp = doc.add_paragraph()
            sp.space_after = Pt(4)
            sr = sp.add_run(', '.join(str(s) for s in skills))
            sr.font.size = Pt(10)

    def _add_experience_entry(self, doc, exp: dict):
        """Add a work experience entry with title, meta, stack, bullets."""
        from docx.shared import Pt, RGBColor

        title = exp.get('title', 'Role')
        company = exp.get('company', 'Company')

        # ── Title line ──
        t_para = doc.add_paragraph()
        t_para.space_before = Pt(2)
        t_para.space_after = Pt(1)

        t_run = t_para.add_run(title)
        t_run.bold = True
        t_run.font.size = Pt(10.5)
        t_run.font.color.rgb = RGBColor(*self.COLOR_PRIMARY)

        sep_run = t_para.add_run(' — ')
        sep_run.font.size = Pt(10.5)

        c_run = t_para.add_run(company)
        c_run.font.size = Pt(10.5)
        c_run.font.color.rgb = RGBColor(*self.COLOR_SECONDARY)

        # ── Duration / Location meta ──
        meta_parts = []
        if exp.get('duration'):
            meta_parts.append(str(exp['duration']))
        elif exp.get('start_date'):
            end = exp.get('end_date', 'Present')
            meta_parts.append(f"{exp['start_date']} — {end}")
        if exp.get('location'):
            meta_parts.append(str(exp['location']))

        if meta_parts:
            m_para = doc.add_paragraph()
            m_para.space_before = Pt(0)
            m_para.space_after = Pt(2)
            m_run = m_para.add_run(' │ '.join(meta_parts))
            m_run.font.size = Pt(9)
            m_run.font.color.rgb = RGBColor(*self.COLOR_META)
            m_run.italic = True

        # ── Stack ──
        if exp.get('stack'):
            stack = exp['stack']
            if isinstance(stack, list):
                stack = ', '.join(str(s) for s in stack)
            sk_para = doc.add_paragraph()
            sk_para.space_before = Pt(0)
            sk_para.space_after = Pt(2)
            sk_label = sk_para.add_run("Stack: ")
            sk_label.bold = True
            sk_label.font.size = Pt(9.5)
            sk_val = sk_para.add_run(str(stack))
            sk_val.font.size = Pt(9.5)

        # ── Description bullets ──
        bullets = self._extract_bullets(exp)
        for b in bullets:
            bp = doc.add_paragraph(style='List Bullet')
            bp.space_before = Pt(0)
            bp.space_after = Pt(1)
            br = bp.add_run(b)
            br.font.size = Pt(10)

        # ── Projects (sub-entries) ──
        if exp.get('projects'):
            proj_items = self._normalize_list(exp['projects'])
            for proj_text in proj_items:
                if proj_text:
                    pp = doc.add_paragraph(style='List Bullet')
                    pp.space_before = Pt(0)
                    pp.space_after = Pt(1)
                    pr = pp.add_run(proj_text)
                    pr.font.size = Pt(10)

        # ── Recognition ──
        if exp.get('recognition'):
            rec_items = self._normalize_list(exp['recognition'])
            for r in rec_items:
                rp = doc.add_paragraph(style='List Bullet')
                rp.space_before = Pt(0)
                rp.space_after = Pt(1)
                rr = rp.add_run(f"★ {r}")
                rr.font.size = Pt(10)
                rr.bold = True
                rr.font.color.rgb = RGBColor(*self.COLOR_ACCENT)

    def _add_project_entry(self, doc, project: dict):
        """Add a standalone project entry."""
        from docx.shared import Pt, RGBColor

        name = project.get('name', 'Project')

        # ── Project name ──
        t_para = doc.add_paragraph()
        t_para.space_before = Pt(2)
        t_para.space_after = Pt(1)
        t_run = t_para.add_run(name)
        t_run.bold = True
        t_run.font.size = Pt(10.5)
        t_run.font.color.rgb = RGBColor(*self.COLOR_PRIMARY)

        # ── Tech stack ──
        tech = project.get('tech', project.get('technologies', project.get('stack', '')))
        if tech:
            if isinstance(tech, list):
                tech = ', '.join(str(t) for t in tech)
            tk_para = doc.add_paragraph()
            tk_para.space_before = Pt(0)
            tk_para.space_after = Pt(2)
            tk_label = tk_para.add_run("Tech: ")
            tk_label.bold = True
            tk_label.font.size = Pt(9.5)
            tk_val = tk_para.add_run(str(tech))
            tk_val.font.size = Pt(9.5)

        # ── Description bullets ──
        bullets = self._extract_bullets(project)
        for b in bullets:
            bp = doc.add_paragraph(style='List Bullet')
            bp.space_before = Pt(0)
            bp.space_after = Pt(1)
            br = bp.add_run(b)
            br.font.size = Pt(10)

        # ── Features ──
        if project.get('features'):
            feat_items = self._normalize_list(project['features'])
            for f in feat_items:
                if f:
                    fp = doc.add_paragraph(style='List Bullet')
                    fp.space_before = Pt(0)
                    fp.space_after = Pt(1)
                    fr = fp.add_run(f)
                    fr.font.size = Pt(10)

    def _add_education_entry(self, doc, edu: dict):
        """Add an education entry."""
        from docx.shared import Pt, RGBColor

        degree = edu.get('degree', 'Degree')
        university = edu.get('university', edu.get('institution', 'University'))

        # ── Degree — University ──
        t_para = doc.add_paragraph()
        t_para.space_before = Pt(2)
        t_para.space_after = Pt(1)

        d_run = t_para.add_run(degree)
        d_run.bold = True
        d_run.font.size = Pt(10.5)
        d_run.font.color.rgb = RGBColor(*self.COLOR_PRIMARY)

        sep_run = t_para.add_run(' — ')
        sep_run.font.size = Pt(10.5)

        u_run = t_para.add_run(university)
        u_run.font.size = Pt(10.5)
        u_run.font.color.rgb = RGBColor(*self.COLOR_SECONDARY)

        # ── Year / CGPA ──
        meta_parts = []
        if edu.get('year'):
            meta_parts.append(str(edu['year']))
        elif edu.get('start_year') and edu.get('end_year'):
            meta_parts.append(f"{edu['start_year']} — {edu['end_year']}")
        if edu.get('cgpa'):
            meta_parts.append(f"CGPA: {edu['cgpa']}")
        if edu.get('percentage'):
            meta_parts.append(f"{edu['percentage']}%")

        if meta_parts:
            m_para = doc.add_paragraph()
            m_para.space_before = Pt(0)
            m_para.space_after = Pt(2)
            m_run = m_para.add_run(' │ '.join(meta_parts))
            m_run.font.size = Pt(9)
            m_run.font.color.rgb = RGBColor(*self.COLOR_META)
            m_run.italic = True

    # ════════════════════════════════════════════════════════════
    #  HTML / PDF GENERATION
    # ════════════════════════════════════════════════════════════

    def _generate_html(self, data: dict) -> str:
        """Generate complete ATS-friendly HTML for PDF conversion."""
        name = self._esc(data.get('name', 'Name'))
        email = self._esc(data.get('email', ''))
        phone = self._esc(data.get('phone', ''))
        location = self._esc(data.get('location', ''))
        linkedin = data.get('linkedin', '')
        github = data.get('github', '')
        summary = self._esc(data.get('summary', ''))

        # ── Contact line ──
        contact_parts = [p for p in [email, phone, location] if p]
        contact_line = ' │ '.join(contact_parts)

        # ── Links line ──
        link_parts = []
        if linkedin:
            url = linkedin if linkedin.startswith('http') else f"https://{linkedin}"
            link_parts.append(f'<a href="{self._esc(url)}">{self._esc(linkedin)}</a>')
        if github:
            url = github if github.startswith('http') else f"https://{github}"
            link_parts.append(f'<a href="{self._esc(url)}">{self._esc(github)}</a>')
        links_line = ' │ '.join(link_parts)

        # ── Build sections ──
        skills_html = self._skills_to_html(data.get('skills', {}))
        exp_html = ''.join(self._experience_to_html(e) for e in data.get('experience', []))
        proj_html = ''.join(self._project_to_html(p) for p in data.get('projects', []))
        edu_html = ''.join(self._education_to_html(e) for e in data.get('education', []))

        cert_html = ''
        for cert in data.get('certifications', []):
            ct = self._cert_to_text(cert)
            if ct:
                cert_html += f'<li>{self._esc(ct)}</li>\n'

        ach_html = ''
        for ach in data.get('achievements', []):
            ach_html += f'<li>{self._esc(str(ach))}</li>\n'

        coding_html = ''
        coding_profiles = data.get('coding_profiles', {})
        if coding_profiles:
            cp_parts = []
            for platform, detail in coding_profiles.items():
                if isinstance(detail, str):
                    cp_parts.append(f"{self._esc(platform)}: {self._esc(detail)}")
                elif isinstance(detail, dict):
                    desc = detail.get('description', detail.get('url', ''))
                    cp_parts.append(f"{self._esc(platform)}: {self._esc(desc)}")
            if cp_parts:
                coding_html = '<p>' + ' │ '.join(cp_parts) + '</p>'

        # ── Assemble full HTML ──
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{name} — Resume</title>
<style>
{self._get_css()}
</style>
</head>
<body>

<div class="header">
    <h1>{name}</h1>
    <div class="contact">{contact_line}</div>
    <div class="links">{links_line}</div>
</div>
"""

        if summary:
            html += f"""
<div class="section">
    <div class="section-title">Professional Summary</div>
    <p>{summary}</p>
</div>
"""

        if skills_html:
            html += f"""
<div class="section">
    <div class="section-title">Technical Skills</div>
    {skills_html}
</div>
"""

        if exp_html:
            html += f"""
<div class="section">
    <div class="section-title">Work Experience</div>
    {exp_html}
</div>
"""

        if proj_html:
            html += f"""
<div class="section">
    <div class="section-title">Projects</div>
    {proj_html}
</div>
"""

        if edu_html:
            html += f"""
<div class="section">
    <div class="section-title">Education</div>
    {edu_html}
</div>
"""

        if cert_html:
            html += f"""
<div class="section">
    <div class="section-title">Certifications</div>
    <ul>{cert_html}</ul>
</div>
"""

        if ach_html:
            html += f"""
<div class="section">
    <div class="section-title">Achievements</div>
    <ul>{ach_html}</ul>
</div>
"""

        if coding_html:
            html += f"""
<div class="section">
    <div class="section-title">Coding Profiles</div>
    {coding_html}
</div>
"""

        html += """
</body>
</html>"""

        return html

    def _get_css(self) -> str:
        """Return CSS stylesheet for HTML resume."""
        return """
    @page {
        size: A4;
        margin: 1.2cm 1.5cm;
    }
    * {
        margin: 0;
        padding: 0;
        box-sizing: border-box;
    }
    body {
        font-family: 'Calibri', 'Segoe UI', 'Helvetica Neue', Arial, sans-serif;
        font-size: 10pt;
        color: #333333;
        line-height: 1.4;
    }
    .header {
        text-align: center;
        margin-bottom: 8px;
        padding-bottom: 8px;
        border-bottom: 1.5px solid #1A1A2E;
    }
    .header h1 {
        font-size: 18pt;
        color: #1A1A2E;
        letter-spacing: 2px;
        margin-bottom: 4px;
        text-transform: uppercase;
    }
    .header .contact {
        font-size: 9pt;
        color: #555555;
        margin-bottom: 2px;
    }
    .header .links {
        font-size: 9pt;
        color: #555555;
    }
    .header .links a {
        color: #2962FF;
        text-decoration: none;
    }
    .section {
        margin-top: 10px;
        margin-bottom: 4px;
    }
    .section-title {
        font-size: 11pt;
        font-weight: bold;
        color: #1A1A2E;
        text-transform: uppercase;
        letter-spacing: 1px;
        border-bottom: 1.5px solid #1A1A2E;
        padding-bottom: 2px;
        margin-bottom: 6px;
    }
    .entry {
        margin-bottom: 8px;
    }
    .entry-header {
        margin-bottom: 2px;
    }
    .entry-title {
        font-size: 10.5pt;
        font-weight: bold;
        color: #1A1A2E;
    }
    .entry-company {
        color: #444444;
        font-weight: normal;
    }
    .entry-meta {
        font-size: 9pt;
        color: #777777;
        font-style: italic;
        margin-bottom: 3px;
    }
    .entry-stack {
        font-size: 9.5pt;
        margin-bottom: 3px;
    }
    .entry-stack strong {
        font-weight: bold;
    }
    ul {
        margin-left: 18px;
        margin-bottom: 4px;
    }
    li {
        margin-bottom: 1px;
        font-size: 10pt;
    }
    .skill-line {
        margin-bottom: 2px;
        font-size: 10pt;
    }
    .recognition {
        color: #2E7D32;
        font-weight: bold;
    }
    p {
        margin-bottom: 4px;
    }
"""

    def _skills_to_html(self, skills) -> str:
        """Convert skills dict/list to HTML."""
        if not skills:
            return ''

        html = ''
        if isinstance(skills, dict):
            for cat, items in skills.items():
                if not items:
                    continue
                items_str = ', '.join(str(s) for s in items) if isinstance(items, list) else str(items)
                html += (
                    f'<p class="skill-line">'
                    f'<strong>{self._esc(cat)}:</strong> {self._esc(items_str)}'
                    f'</p>\n'
                )
        elif isinstance(skills, list):
            html = f'<p>{self._esc(", ".join(str(s) for s in skills))}</p>'

        return html

    def _experience_to_html(self, exp: dict) -> str:
        """Convert experience entry to HTML."""
        title = self._esc(exp.get('title', 'Role'))
        company = self._esc(exp.get('company', 'Company'))

        # Meta
        meta_parts = []
        if exp.get('duration'):
            meta_parts.append(self._esc(str(exp['duration'])))
        elif exp.get('start_date'):
            end = exp.get('end_date', 'Present')
            meta_parts.append(f"{self._esc(exp['start_date'])} — {self._esc(end)}")
        if exp.get('location'):
            meta_parts.append(self._esc(str(exp['location'])))
        meta_line = ' │ '.join(meta_parts)

        # Stack
        stack_html = ''
        if exp.get('stack'):
            stack = exp['stack']
            if isinstance(stack, list):
                stack = ', '.join(str(s) for s in stack)
            stack_html = (
                f'<div class="entry-stack">'
                f'<strong>Stack:</strong> {self._esc(str(stack))}'
                f'</div>'
            )

        # Bullets
        all_bullets = self._extract_bullets(exp)

        # Projects
        if exp.get('projects'):
            proj_items = self._normalize_list(exp['projects'])
            all_bullets.extend(proj_items)

        # Recognition
        recognition_html = ''
        if exp.get('recognition'):
            rec_items = self._normalize_list(exp['recognition'])
            for r in rec_items:
                all_bullets.append(f'★ {r}')

        bullets_html = ''
        if all_bullets:
            bullets_html = '<ul>\n'
            for b in all_bullets:
                if b.startswith('★'):
                    bullets_html += f'<li class="recognition">{self._esc(b)}</li>\n'
                else:
                    bullets_html += f'<li>{self._esc(b)}</li>\n'
            bullets_html += '</ul>\n'

        return f"""
<div class="entry">
    <div class="entry-header">
        <span class="entry-title">{title} — <span class="entry-company">{company}</span></span>
    </div>
    <div class="entry-meta">{meta_line}</div>
    {stack_html}
    {bullets_html}
</div>
"""

    def _project_to_html(self, project: dict) -> str:
        """Convert project entry to HTML."""
        name = self._esc(project.get('name', 'Project'))

        tech = project.get('tech', project.get('technologies', project.get('stack', '')))
        tech_html = ''
        if tech:
            if isinstance(tech, list):
                tech = ', '.join(str(t) for t in tech)
            tech_html = (
                f'<div class="entry-stack">'
                f'<strong>Tech:</strong> {self._esc(str(tech))}'
                f'</div>'
            )

        all_bullets = self._extract_bullets(project)
        if project.get('features'):
            feat_items = self._normalize_list(project['features'])
            all_bullets.extend(feat_items)

        bullets_html = ''
        if all_bullets:
            bullets_html = '<ul>\n'
            for b in all_bullets:
                bullets_html += f'<li>{self._esc(b)}</li>\n'
            bullets_html += '</ul>\n'

        return f"""
<div class="entry">
    <div class="entry-title">{name}</div>
    {tech_html}
    {bullets_html}
</div>
"""

    def _education_to_html(self, edu: dict) -> str:
        """Convert education entry to HTML."""
        degree = self._esc(edu.get('degree', 'Degree'))
        university = self._esc(edu.get('university', edu.get('institution', 'University')))

        meta_parts = []
        if edu.get('year'):
            meta_parts.append(str(edu['year']))
        elif edu.get('start_year') and edu.get('end_year'):
            meta_parts.append(f"{edu['start_year']} — {edu['end_year']}")
        if edu.get('cgpa'):
            meta_parts.append(f"CGPA: {edu['cgpa']}")
        if edu.get('percentage'):
            meta_parts.append(f"{edu['percentage']}%")
        meta_line = ' │ '.join(meta_parts)

        return f"""
<div class="entry">
    <div class="entry-title">{degree} — <span class="entry-company">{university}</span></div>
    <div class="entry-meta">{meta_line}</div>
</div>
"""

    # ════════════════════════════════════════════════════════════
    #  UTILITY HELPERS
    # ════════════════════════════════════════════════════════════

    def _to_dict(self, resume_data) -> dict:
        """Convert ResumeData dataclass or any object to dict."""
        if isinstance(resume_data, dict):
            return resume_data

        # dataclass
        try:
            from dataclasses import asdict
            return asdict(resume_data)
        except (TypeError, ImportError):
            pass

        # Manual attribute extraction
        fields = [
            'name', 'email', 'phone', 'location', 'linkedin', 'github',
            'summary', 'education', 'experience', 'projects', 'skills',
            'certifications', 'achievements', 'coding_profiles',
            '_target_company', '_target_title',
        ]
        result = {}
        for f in fields:
            val = getattr(resume_data, f, None)
            if val is not None:
                result[f] = val
        if not result:
            logger.error(f"Cannot convert resume_data of type {type(resume_data)} to dict")
        return result

    @staticmethod
    def _sanitize_filename(name: str) -> str:
        """Sanitize string for filesystem use."""
        sanitized = re.sub(r'[^\w\s-]', '', str(name))
        sanitized = re.sub(r'\s+', '_', sanitized.strip())
        return sanitized[:50] or 'Unknown'

    @staticmethod
    def _esc(text) -> str:
        """HTML-escape a string."""
        if not text:
            return ''
        text = str(text)
        return (
            text.replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#39;')
        )

    @staticmethod
    def _extract_bullets(entry: dict) -> List[str]:
        """Extract bullet points from description field."""
        desc = entry.get('description', '')
        if not desc:
            return []

        if isinstance(desc, str):
            raw = [b.strip() for b in desc.split('\n') if b.strip()]
        elif isinstance(desc, list):
            raw = [str(b).strip() for b in desc if b]
        else:
            raw = [str(desc)]

        # Strip leading bullet characters
        return [b.lstrip('•●○◦-–— ').strip() for b in raw if b.lstrip('•●○◦-–— ').strip()]

    @staticmethod
    def _normalize_list(value) -> List[str]:
        """
        Normalize a value to a flat list of strings.
        Handles str, list[str], list[dict], etc.
        """
        if not value:
            return []

        if isinstance(value, str):
            return [p.strip() for p in value.split(',') if p.strip()]

        if isinstance(value, list):
            result = []
            for item in value:
                if isinstance(item, dict):
                    name = item.get('name', '')
                    desc = item.get('description', '')
                    text = f"{name}: {desc}" if desc else name
                    if text:
                        result.append(text.strip())
                elif item:
                    result.append(str(item).strip())
            return result

        return [str(value)]

    @staticmethod
    def _cert_to_text(cert) -> str:
        """Convert a certification entry to text string."""
        if isinstance(cert, dict):
            text = cert.get('name', str(cert))
            if cert.get('issuer'):
                text += f" — {cert['issuer']}"
            if cert.get('year'):
                text += f" ({cert['year']})"
            return text
        elif cert:
            return str(cert)
        return ''


# ════════════════════════════════════════════════════════════════
#  TEST BLOCK
# ════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import json

    print("=" * 65)
    print("  RESUME BUILDER — Test Suite")
    print("=" * 65)

    # ── Try to load real resume data ──
    try:
        from profile.resume_data import get_base_resume, resume_to_dict
        base_resume = get_base_resume()
        test_data = resume_to_dict(base_resume)
        print("\n✅ Loaded resume data from profile/resume_data.py")
    except Exception as e:
        print(f"\n⚠️  Could not load profile/resume_data.py ({e})")
        print("   Using built-in test data instead.\n")

        # Fallback: minimal test data matching Piyush's profile
        test_data = {
            'name': 'Piyush Kashyap',
            'email': 'piyushkashyap3247@gmail.com',
            'phone': '+91 73107 03247',
            'location': 'Rishikesh, Uttarakhand, India',
            'linkedin': 'linkedin.com/in/piyush-kashyap731',
            'github': 'github.com/Piyush731',
            'summary': (
                'Full Stack Developer L1 with end-to-end responsibility of 10+ production '
                'applications across fintech, ERP, edtech, and CRM domains. Sole developer '
                'on each project — handling client requirements, database design, frontend, '
                'and backend independently. Shipped multi-tenant systems with 57+ DB tables, '
                'real-time WebSocket feeds, and third-party integrations (META WhatsApp, '
                'Razorpay, RTO API, IOTrades). Published app on Google Play Store.'
            ),
            'skills': {
                'Languages': ['JavaScript (ES6+)', 'Java', 'Python', 'SQL'],
                'Frontend': ['Vue.js', 'Nuxt.js', 'React.js', 'Vuetify', 'Tailwind CSS'],
                'Backend': ['Node.js', 'Express.js', 'Spring Boot', 'REST APIs', 'WebSockets', 'Microservices'],
                'Databases': ['MySQL', 'MongoDB', 'PostgreSQL', 'Redis'],
                'Integrations': ['Razorpay', 'META WhatsApp API', 'RTO API', 'IOTrades API'],
                'Tools': ['Git', 'Docker', 'Kafka', 'JWT', 'Playwright'],
            },
            'experience': [
                {
                    'title': 'Full Stack Developer L1',
                    'company': 'Site Guru Pvt Ltd',
                    'duration': 'Aug 2024 — Present',
                    'location': 'Remote',
                    'stack': 'Vue.js, Nuxt.js, Node.js, MySQL, Vuetify, REST APIs, WebSockets',
                    'description': [
                        'Sole developer on 10+ production applications across fintech, ERP, edtech, and CRM domains',
                        'Built BizHub ERP — multi-tenant system with 57 DB tables, 5 user types, full business flow',
                        'Developed My RTO Expert handling 1000+ daily requests with META WhatsApp API integration',
                        'Created Rudra Fintech — investment platform with interest/TDS calculation and audit logging',
                        'Built FX Prime Trading platform with IOTrades API, MT5 integration, WebSocket price feeds',
                        'Shipped Dheeranet ISP CRM with hierarchical zone management and IP allocation',
                    ],
                    'projects': [
                        'BizHub ERP', 'My RTO Expert', 'Rudra Fintech',
                        'FX Prime Trading', 'Dheeranet ISP CRM',
                        'SB Flying Services', 'TutorsUp (Play Store)', 'SoloWash',
                    ],
                    'recognition': [
                        'Top performer — given production deployment authority',
                        'Mentor to new team members',
                    ],
                },
                {
                    'title': 'Salesforce Developer Intern',
                    'company': 'SmartBridge',
                    'duration': 'July — Sept 2024',
                    'description': [
                        'Built Apex triggers and batch classes for automated lead processing',
                        'Developed Lightning Web Components (LWC) for custom UI',
                        'Implemented automated lead assignment reducing manual work by 30%',
                    ],
                },
            ],
            'projects': [
                {
                    'name': 'Collaborative Workspace',
                    'tech': ['React', 'Node.js', 'MongoDB', 'Socket.io', 'Gitea API'],
                    'description': [
                        'Real-time collaborative platform with RBAC for 50+ users',
                        'Integrated Gitea API for version control within the workspace',
                    ],
                },
                {
                    'name': 'Invoice Microservice',
                    'tech': ['Java', 'Spring Boot', 'PostgreSQL', 'Kafka', 'Docker'],
                    'description': [
                        'Event-driven microservice architecture for invoice processing',
                        'Automated PDF export with configurable templates',
                    ],
                },
                {
                    'name': 'CareerCraft AI Resume Analyzer',
                    'tech': ['Python', 'Gemini API', 'Streamlit'],
                    'description': [
                        'NLP-powered resume analysis with JD match scoring',
                        'Provides keyword gap analysis and optimization suggestions',
                    ],
                },
            ],
            'education': [
                {
                    'degree': 'B.Tech Computer Science',
                    'university': 'Graphic Era Hill University',
                    'year': '2021 — 2025',
                    'cgpa': '7.79/10',
                },
            ],
            'certifications': [
                {'name': 'Full Stack Java Developer', 'issuer': 'Udemy', 'year': '2024'},
                {'name': 'Agile Project Management', 'issuer': 'Udemy', 'year': '2024'},
            ],
            'achievements': [
                'Solved 100+ problems on LeetCode/GFG (Arrays, Trees, Graphs, DP)',
                'Published TutorsUp app on Google Play Store',
                'Top performer recognition at Site Guru — production deployment authority',
            ],
        }

    # ── Initialize builder ──
    builder = ResumeBuilder()

    # ── Test 1: DOCX generation ──
    print("\n─── Test 1: Build DOCX ─────────────────────────────────")
    try:
        docx_path = builder.build_docx(
            test_data,
            company='TestCompany',
            title='SDE-1',
        )
        file_size = os.path.getsize(docx_path)
        print(f"  ✅ DOCX generated: {docx_path}")
        print(f"     Size: {file_size:,} bytes ({file_size / 1024:.1f} KB)")
    except Exception as e:
        print(f"  ❌ DOCX failed: {e}")
        import traceback
        traceback.print_exc()

    # ── Test 2: PDF generation ──
    print("\n─── Test 2: Build PDF ──────────────────────────────────")
    try:
        pdf_path = builder.build_pdf(
            test_data,
            company='TestCompany',
            title='Backend-Developer',
        )
        file_size = os.path.getsize(pdf_path)
        ext = os.path.splitext(pdf_path)[1]
        label = 'PDF' if ext == '.pdf' else 'HTML (weasyprint not available)'
        print(f"  ✅ {label} generated: {pdf_path}")
        print(f"     Size: {file_size:,} bytes ({file_size / 1024:.1f} KB)")
    except Exception as e:
        print(f"  ❌ PDF failed: {e}")
        import traceback
        traceback.print_exc()

    # ── Test 3: Build both ──
    print("\n─── Test 3: Build Both ─────────────────────────────────")
    try:
        results = builder.build_both(
            test_data,
            company='Razorpay',
            title='Full-Stack-Developer',
        )
        for fmt, path in results.items():
            if path:
                size = os.path.getsize(path)
                print(f"  ✅ {fmt.upper()}: {path} ({size / 1024:.1f} KB)")
            else:
                print(f"  ⚠️  {fmt.upper()}: Not generated")
    except Exception as e:
        print(f"  ❌ Build both failed: {e}")
        import traceback
        traceback.print_exc()

    # ── Test 4: Output path generation ──
    print("\n─── Test 4: Output Path ────────────────────────────────")
    path1 = builder.get_output_path("Google India", "SDE-1 Backend")
    path2 = builder.get_output_path("Razorpay", "Full Stack Developer")
    print(f"  Path 1: {path1}")
    print(f"  Path 2: {path2}")

    # ── Test 5: Filename sanitization ──
    print("\n─── Test 5: Filename Sanitization ──────────────────────")
    test_names = [
        "Google (India) Pvt. Ltd.",
        "SDE-1 / Backend Developer",
        "   Spaces   Everywhere   ",
        "Special!@#$%^&*()Characters",
        "",
    ]
    for name in test_names:
        sanitized = ResumeBuilder._sanitize_filename(name)
        print(f"  '{name}' → '{sanitized}'")

    # ── Test 6: Template placeholder (dry run) ──
    print("\n─── Test 6: Template Mode (info) ───────────────────────")
    template_path = os.path.join(BASE_DIR, 'profile', 'templates', 'resume_base.docx')
    if os.path.exists(template_path):
        print(f"  ✅ Template found: {template_path}")
        try:
            docx_t = builder.build_docx(test_data, template_path=template_path,
                                         company='TemplateCo', title='TemplateRole')
            print(f"  ✅ Template-based DOCX: {docx_t}")
        except Exception as e:
            print(f"  ⚠️  Template build error: {e}")
    else:
        print(f"  ℹ️  No template at {template_path}")
        print("     Create a .docx with {{NAME}}, {{SUMMARY}}, etc. placeholders")
        print("     Supported: {{NAME}} {{EMAIL}} {{PHONE}} {{LOCATION}}")
        print("                {{LINKEDIN}} {{GITHUB}} {{SUMMARY}} {{SKILLS}}")
        print("                {{EXPERIENCE}} {{PROJECTS}} {{EDUCATION}}")
        print("                {{CERTIFICATIONS}} {{ACHIEVEMENTS}}")

    # ── Summary ──
    print("\n" + "=" * 65)
    output_files = list(Path(RESUME_OUTPUT_DIR).rglob('*'))
    output_files = [f for f in output_files if f.is_file()]
    print(f"  Total files generated: {len(output_files)}")
    print(f"  Output directory: {RESUME_OUTPUT_DIR}")
    print("=" * 65)
    print("  ✅ resume/builder.py — All tests complete")
    print("=" * 65)