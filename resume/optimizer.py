"""
resume/optimizer.py — ATS Resume Optimizer
═══════════════════════════════════════════════════════════════
Scores resumes against job descriptions, identifies keyword gaps,
optimizes section ordering, and checks DOCX formatting for ATS
compliance. Pure algorithmic — no LLM required (optional enhance).

Interface:
  ATSOptimizer()
  ATSOptimizer.optimize(resume_data, job) → ResumeData
  ATSOptimizer.score(resume_data, job) → dict
  ATSOptimizer.check_formatting(docx_path) → dict
  ATSOptimizer.extract_keywords(text) → dict
  ATSOptimizer.keyword_gap(resume_data, job) → dict

Dependencies: config.py, core/logger.py, profile/resume_data.py
"""

import os
import re
import sys
import math
import string
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, List, Set, Tuple, Any, Union
from collections import Counter

# ─── Project root on path ──────────────────────────────────────
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import RESUME_CONFIG, MATCH_CONFIG, BASE_DIR
from core.logger import get_logger

logger = get_logger("resume.optimizer")


# ════════════════════════════════════════════════════════════════
#  KNOWN TECH SKILLS — exhaustive lookup for extraction
# ════════════════════════════════════════════════════════════════

# Canonical skill names (lowercase → display name)
KNOWN_SKILLS = {
    # Languages
    'javascript': 'JavaScript', 'js': 'JavaScript', 'es6': 'JavaScript (ES6+)',
    'typescript': 'TypeScript', 'ts': 'TypeScript',
    'python': 'Python', 'python3': 'Python',
    'java': 'Java', 'core java': 'Java',
    'c++': 'C++', 'cpp': 'C++', 'c/c++': 'C/C++',
    'c#': 'C#', 'csharp': 'C#',
    'go': 'Go', 'golang': 'Go',
    'rust': 'Rust', 'ruby': 'Ruby', 'php': 'PHP',
    'kotlin': 'Kotlin', 'swift': 'Swift', 'scala': 'Scala',
    'r': 'R', 'matlab': 'MATLAB', 'perl': 'Perl',
    'dart': 'Dart', 'elixir': 'Elixir', 'haskell': 'Haskell',
    'sql': 'SQL', 'plsql': 'PL/SQL', 'pl/sql': 'PL/SQL',
    'bash': 'Bash', 'shell': 'Shell Scripting',
    'html': 'HTML', 'html5': 'HTML5',
    'css': 'CSS', 'css3': 'CSS3', 'sass': 'SASS', 'scss': 'SCSS', 'less': 'LESS',

    # Frontend Frameworks
    'react': 'React.js', 'react.js': 'React.js', 'reactjs': 'React.js',
    'vue': 'Vue.js', 'vue.js': 'Vue.js', 'vuejs': 'Vue.js',
    'angular': 'Angular', 'angularjs': 'AngularJS',
    'next.js': 'Next.js', 'nextjs': 'Next.js', 'next': 'Next.js',
    'nuxt': 'Nuxt.js', 'nuxt.js': 'Nuxt.js', 'nuxtjs': 'Nuxt.js',
    'svelte': 'Svelte', 'sveltekit': 'SvelteKit',
    'gatsby': 'Gatsby', 'remix': 'Remix',
    'jquery': 'jQuery',

    # UI Libraries
    'tailwind': 'Tailwind CSS', 'tailwindcss': 'Tailwind CSS', 'tailwind css': 'Tailwind CSS',
    'bootstrap': 'Bootstrap', 'material ui': 'Material UI', 'mui': 'Material UI',
    'vuetify': 'Vuetify', 'ant design': 'Ant Design', 'antd': 'Ant Design',
    'chakra ui': 'Chakra UI', 'shadcn': 'shadcn/ui',

    # Backend
    'node': 'Node.js', 'node.js': 'Node.js', 'nodejs': 'Node.js',
    'express': 'Express.js', 'express.js': 'Express.js', 'expressjs': 'Express.js',
    'fastify': 'Fastify', 'koa': 'Koa', 'nestjs': 'NestJS', 'nest.js': 'NestJS',
    'spring': 'Spring', 'spring boot': 'Spring Boot', 'springboot': 'Spring Boot',
    'spring mvc': 'Spring MVC', 'spring cloud': 'Spring Cloud',
    'hibernate': 'Hibernate', 'jpa': 'JPA',
    'django': 'Django', 'flask': 'Flask', 'fastapi': 'FastAPI',
    'rails': 'Ruby on Rails', 'ruby on rails': 'Ruby on Rails',
    'laravel': 'Laravel', '.net': '.NET', 'asp.net': 'ASP.NET',
    'graphql': 'GraphQL', 'grpc': 'gRPC',

    # Databases
    'mysql': 'MySQL', 'postgresql': 'PostgreSQL', 'postgres': 'PostgreSQL',
    'mongodb': 'MongoDB', 'mongo': 'MongoDB',
    'redis': 'Redis', 'elasticsearch': 'Elasticsearch',
    'sqlite': 'SQLite', 'oracle': 'Oracle DB',
    'mssql': 'MS SQL Server', 'sql server': 'MS SQL Server',
    'dynamodb': 'DynamoDB', 'cassandra': 'Cassandra',
    'firebase': 'Firebase', 'firestore': 'Firestore',
    'supabase': 'Supabase', 'cockroachdb': 'CockroachDB',
    'neo4j': 'Neo4j', 'couchdb': 'CouchDB',
    'mariadb': 'MariaDB', 'memcached': 'Memcached',

    # APIs & Protocols
    'rest': 'REST APIs', 'rest api': 'REST APIs', 'rest apis': 'REST APIs',
    'restful': 'RESTful APIs', 'websocket': 'WebSockets', 'websockets': 'WebSockets',
    'socket.io': 'Socket.io', 'socketio': 'Socket.io',
    'soap': 'SOAP', 'webhook': 'Webhooks', 'webhooks': 'Webhooks',

    # DevOps & Cloud
    'docker': 'Docker', 'kubernetes': 'Kubernetes', 'k8s': 'Kubernetes',
    'aws': 'AWS', 'amazon web services': 'AWS',
    'gcp': 'Google Cloud', 'google cloud': 'Google Cloud',
    'azure': 'Microsoft Azure', 'heroku': 'Heroku',
    'vercel': 'Vercel', 'netlify': 'Netlify', 'digitalocean': 'DigitalOcean',
    'terraform': 'Terraform', 'ansible': 'Ansible',
    'jenkins': 'Jenkins', 'github actions': 'GitHub Actions',
    'gitlab ci': 'GitLab CI', 'ci/cd': 'CI/CD', 'cicd': 'CI/CD',
    'nginx': 'Nginx', 'apache': 'Apache', 'pm2': 'PM2',
    'linux': 'Linux', 'ubuntu': 'Ubuntu',

    # Message Queues & Streaming
    'kafka': 'Kafka', 'apache kafka': 'Kafka',
    'rabbitmq': 'RabbitMQ', 'sqs': 'Amazon SQS',
    'redis pub/sub': 'Redis Pub/Sub', 'celery': 'Celery',
    'bull': 'Bull Queue', 'zeromq': 'ZeroMQ',

    # Architecture
    'microservices': 'Microservices', 'monolith': 'Monolithic',
    'event-driven': 'Event-Driven Architecture',
    'serverless': 'Serverless', 'soa': 'SOA',
    'mvc': 'MVC', 'mvvm': 'MVVM',

    # Auth & Security
    'jwt': 'JWT', 'oauth': 'OAuth', 'oauth2': 'OAuth 2.0',
    'sso': 'SSO', 'saml': 'SAML', 'ldap': 'LDAP',
    'bcrypt': 'bcrypt', 'cors': 'CORS', 'https': 'HTTPS',

    # Testing
    'jest': 'Jest', 'mocha': 'Mocha', 'chai': 'Chai',
    'pytest': 'pytest', 'junit': 'JUnit', 'selenium': 'Selenium',
    'cypress': 'Cypress', 'playwright': 'Playwright',
    'postman': 'Postman', 'swagger': 'Swagger',

    # VCS & Tools
    'git': 'Git', 'github': 'GitHub', 'gitlab': 'GitLab',
    'bitbucket': 'Bitbucket', 'svn': 'SVN',
    'jira': 'Jira', 'confluence': 'Confluence',
    'trello': 'Trello', 'slack': 'Slack',
    'figma': 'Figma', 'vscode': 'VS Code',

    # Data / ML
    'pandas': 'Pandas', 'numpy': 'NumPy', 'scikit-learn': 'scikit-learn',
    'tensorflow': 'TensorFlow', 'pytorch': 'PyTorch', 'keras': 'Keras',
    'opencv': 'OpenCV', 'nltk': 'NLTK', 'spacy': 'spaCy',
    'hadoop': 'Hadoop', 'spark': 'Apache Spark', 'airflow': 'Apache Airflow',
    'tableau': 'Tableau', 'power bi': 'Power BI',

    # Mobile
    'react native': 'React Native', 'flutter': 'Flutter',
    'android': 'Android', 'ios': 'iOS',
    'swiftui': 'SwiftUI', 'jetpack compose': 'Jetpack Compose',

    # CMS / Platforms
    'wordpress': 'WordPress', 'shopify': 'Shopify',
    'salesforce': 'Salesforce', 'apex': 'Apex (Salesforce)',
    'lwc': 'Lightning Web Components',

    # Third-party / Integrations
    'razorpay': 'Razorpay', 'stripe': 'Stripe', 'paypal': 'PayPal',
    'twilio': 'Twilio', 'sendgrid': 'SendGrid',
    'whatsapp api': 'WhatsApp API', 'meta whatsapp api': 'META WhatsApp API',

    # Concepts
    'agile': 'Agile', 'scrum': 'Scrum', 'kanban': 'Kanban',
    'tdd': 'TDD', 'bdd': 'BDD',
    'oop': 'OOP', 'functional programming': 'Functional Programming',
    'design patterns': 'Design Patterns', 'solid': 'SOLID Principles',
    'data structures': 'Data Structures', 'algorithms': 'Algorithms',
    'dsa': 'DSA', 'system design': 'System Design',
    'multi-tenant': 'Multi-Tenant Architecture',
}

# Multi-word skills that need exact or phrase matching
MULTI_WORD_SKILLS = sorted(
    [k for k in KNOWN_SKILLS if ' ' in k or '.' in k or '/' in k],
    key=len,
    reverse=True,  # longest first for greedy matching
)

# ATS-critical sections in recommended order
ATS_SECTION_ORDER = [
    'summary', 'skills', 'experience', 'projects',
    'education', 'certifications', 'achievements',
]

# Common ATS section header aliases
ATS_SECTION_ALIASES = {
    'summary': ['professional summary', 'summary', 'objective', 'profile', 'about'],
    'skills': ['technical skills', 'skills', 'core competencies', 'technologies'],
    'experience': ['work experience', 'experience', 'professional experience', 'employment'],
    'projects': ['projects', 'key projects', 'personal projects'],
    'education': ['education', 'academic background', 'qualifications'],
    'certifications': ['certifications', 'certificates', 'licenses'],
    'achievements': ['achievements', 'awards', 'recognition', 'accomplishments'],
}


# ════════════════════════════════════════════════════════════════
#  ATS OPTIMIZER
# ════════════════════════════════════════════════════════════════

class ATSOptimizer:
    """
    ATS (Applicant Tracking System) resume optimizer.
    
    - Scores resume vs job description keyword match
    - Identifies missing keywords (gap analysis)
    - Optimizes section ordering for ATS priority
    - Checks DOCX formatting compliance
    - Pure algorithmic — no LLM calls required
    """

    # ── Scoring weights ──
    WEIGHT_KEYWORD_MATCH = 0.40       # How many JD keywords are in resume
    WEIGHT_SKILL_DENSITY = 0.15       # Keyword density in skills section
    WEIGHT_SECTION_ORDER = 0.10       # Correct section ordering
    WEIGHT_SUMMARY_KEYWORDS = 0.10    # Keywords in summary section
    WEIGHT_EXPERIENCE_KEYWORDS = 0.15 # Keywords in experience bullets
    WEIGHT_FORMATTING = 0.10          # ATS-safe formatting

    # ── Formatting rules ──
    ATS_SAFE_FONTS = {
        'calibri', 'arial', 'helvetica', 'times new roman',
        'georgia', 'garamond', 'cambria', 'verdana', 'tahoma',
        'trebuchet ms', 'book antiqua', 'palatino linotype',
    }
    MAX_PAGES = 2
    IDEAL_WORD_COUNT = (400, 800)
    MIN_BULLETS_PER_EXP = 3
    MAX_BULLETS_PER_EXP = 8
    IDEAL_BULLET_WORDS = (8, 25)

    def __init__(self):
        """Initialize optimizer."""
        logger.info("ATSOptimizer initialized")

    # ════════════════════════════════════════════════════════════
    #  PUBLIC: score()
    # ════════════════════════════════════════════════════════════

    def score(self, resume_data, job: dict) -> dict:
        """
        Score resume against a job description for ATS compatibility.
        
        Args:
            resume_data: ResumeData dataclass or dict
            job: Job dict with at least 'description' and optionally
                 'title', 'skills', 'company'
        
        Returns:
            dict: {
                score: 0-100,
                grade: 'A'/'B'/'C'/'D'/'F',
                keyword_score: 0-100,
                density_score: 0-100,
                section_score: 0-100,
                summary_score: 0-100,
                experience_score: 0-100,
                formatting_score: 0-100,
                matched_keywords: [],
                missing_keywords: [],
                total_jd_keywords: int,
                issues: [],
                suggestions: [],
                details: {}
            }
        """
        data = self._to_dict(resume_data)
        jd_text = self._get_jd_text(job)

        if not jd_text:
            logger.warning("No job description provided — returning default score")
            return self._empty_score("No job description available")

        # ── Extract keywords from JD ──
        jd_keywords = self.extract_keywords(jd_text)
        all_jd_skills = set()
        all_jd_skills.update(jd_keywords.get('required_skills', []))
        all_jd_skills.update(jd_keywords.get('preferred_skills', []))
        all_jd_skills.update(jd_keywords.get('tools', []))

        if not all_jd_skills:
            logger.warning("No skills extracted from JD")
            return self._empty_score("Could not extract keywords from job description")

        # ── Extract keywords from resume ──
        resume_text = self._resume_to_text(data)
        resume_keywords = self.extract_keywords(resume_text)
        all_resume_skills = set()
        all_resume_skills.update(resume_keywords.get('required_skills', []))
        all_resume_skills.update(resume_keywords.get('preferred_skills', []))
        all_resume_skills.update(resume_keywords.get('tools', []))

        # Also add explicitly listed skills
        explicit_skills = self._get_explicit_skills(data)
        all_resume_skills.update(explicit_skills)

        # ── Normalize for comparison ──
        jd_normalized = {self._normalize_skill(s) for s in all_jd_skills}
        resume_normalized = {self._normalize_skill(s) for s in all_resume_skills}

        matched_normalized = jd_normalized & resume_normalized
        missing_normalized = jd_normalized - resume_normalized

        # Map back to display names
        jd_display = {self._normalize_skill(s): s for s in all_jd_skills}
        resume_display = {self._normalize_skill(s): s for s in all_resume_skills}

        matched_display = [jd_display.get(n, n) for n in matched_normalized]
        missing_display = [jd_display.get(n, n) for n in missing_normalized]

        issues = []
        suggestions = []

        # ═══════ Sub-scores ═══════

        # 1) Keyword match score (40% weight)
        if jd_normalized:
            keyword_pct = len(matched_normalized) / len(jd_normalized)
        else:
            keyword_pct = 0
        keyword_score = min(100, keyword_pct * 100)

        if keyword_score < 50:
            issues.append(f"Low keyword match: only {len(matched_normalized)}/{len(jd_normalized)} JD keywords found")
            suggestions.append(f"Add missing skills to your resume: {', '.join(missing_display[:8])}")
        elif keyword_score < 75:
            suggestions.append(f"Consider adding: {', '.join(missing_display[:5])}")

        # 2) Skill density score (15% weight)
        skills_section_text = self._get_section_text(data, 'skills')
        density_score = self._calc_density_score(skills_section_text, jd_normalized)

        if density_score < 50:
            issues.append("Skills section has low keyword density vs JD requirements")
            suggestions.append("Move most relevant skills to the top of your skills section")

        # 3) Section ordering score (10% weight)
        section_score = self._calc_section_order_score(data)

        if section_score < 70:
            issues.append("Section ordering not optimal for ATS")
            suggestions.append("Recommended order: Summary → Skills → Experience → Projects → Education")

        # 4) Summary keywords score (10% weight)
        summary_text = data.get('summary', '')
        summary_score = self._calc_text_keyword_score(summary_text, jd_normalized)

        if summary_score < 40:
            issues.append("Summary section lacks JD keywords")
            suggestions.append("Weave 3-5 key technologies from the JD into your summary")

        # 5) Experience keywords score (15% weight)
        exp_text = self._get_section_text(data, 'experience')
        experience_score = self._calc_text_keyword_score(exp_text, jd_normalized)

        if experience_score < 40:
            issues.append("Experience bullets lack JD keywords")
            suggestions.append("Add specific technologies used in each role's bullet points")

        # 6) Formatting score (10% weight)
        formatting_score, fmt_issues, fmt_suggestions = self._calc_formatting_score(data)
        issues.extend(fmt_issues)
        suggestions.extend(fmt_suggestions)

        # ═══════ Weighted total ═══════
        total_score = (
            keyword_score * self.WEIGHT_KEYWORD_MATCH +
            density_score * self.WEIGHT_SKILL_DENSITY +
            section_score * self.WEIGHT_SECTION_ORDER +
            summary_score * self.WEIGHT_SUMMARY_KEYWORDS +
            experience_score * self.WEIGHT_EXPERIENCE_KEYWORDS +
            formatting_score * self.WEIGHT_FORMATTING
        )

        total_score = min(100, max(0, round(total_score)))

        # ── Grade ──
        if total_score >= 85:
            grade = 'A'
        elif total_score >= 70:
            grade = 'B'
        elif total_score >= 55:
            grade = 'C'
        elif total_score >= 40:
            grade = 'D'
        else:
            grade = 'F'

        result = {
            'score': total_score,
            'grade': grade,
            'keyword_score': round(keyword_score),
            'density_score': round(density_score),
            'section_score': round(section_score),
            'summary_score': round(summary_score),
            'experience_score': round(experience_score),
            'formatting_score': round(formatting_score),
            'matched_keywords': sorted(matched_display),
            'missing_keywords': sorted(missing_display),
            'total_jd_keywords': len(jd_normalized),
            'match_ratio': f"{len(matched_normalized)}/{len(jd_normalized)}",
            'issues': issues,
            'suggestions': suggestions,
            'details': {
                'jd_keywords_extracted': sorted(all_jd_skills),
                'resume_keywords_extracted': sorted(all_resume_skills),
                'resume_word_count': len(resume_text.split()),
                'job_title': job.get('title', 'Unknown'),
                'job_company': job.get('company', 'Unknown'),
            },
        }

        logger.info(
            f"ATS Score: {total_score}/100 ({grade}) — "
            f"matched {len(matched_normalized)}/{len(jd_normalized)} keywords"
        )
        return result

    # ════════════════════════════════════════════════════════════
    #  PUBLIC: optimize()
    # ════════════════════════════════════════════════════════════

    def optimize(self, resume_data, job: dict):
        """
        Optimize resume data for ATS compatibility against a job.
        
        NEVER fabricates experience — only:
        - Reorders skills to prioritize JD matches
        - Reorders experience bullets to highlight relevant ones
        - Ensures proper section ordering
        - Injects matched keywords naturally into summary
        - Moves matching projects higher
        
        Args:
            resume_data: ResumeData or dict
            job: Job dict with description
            
        Returns:
            dict: Optimized resume data (same structure, reordered)
        """
        import copy

        data = self._to_dict(resume_data)
        optimized = copy.deepcopy(data)
        jd_text = self._get_jd_text(job)

        if not jd_text:
            logger.warning("No JD to optimize against — returning unchanged")
            return optimized

        jd_keywords = self.extract_keywords(jd_text)
        all_jd_skills = set()
        all_jd_skills.update(jd_keywords.get('required_skills', []))
        all_jd_skills.update(jd_keywords.get('preferred_skills', []))
        all_jd_skills.update(jd_keywords.get('tools', []))

        jd_normalized = {self._normalize_skill(s) for s in all_jd_skills}

        changes = []

        # ── 1) Reorder skills: matching skills first ──
        optimized['skills'] = self._reorder_skills(
            optimized.get('skills', {}), jd_normalized
        )
        changes.append("Reordered skills — JD matches first")

        # ── 2) Reorder experience bullets ──
        if optimized.get('experience'):
            for i, exp in enumerate(optimized['experience']):
                optimized['experience'][i] = self._reorder_bullets(
                    exp, jd_normalized
                )
            changes.append("Reordered experience bullets by JD relevance")

        # ── 3) Reorder projects: matching tech stack first ──
        if optimized.get('projects'):
            optimized['projects'] = self._reorder_projects(
                optimized['projects'], jd_normalized
            )
            changes.append("Reordered projects by JD relevance")

        # ── 4) Enhance summary with JD keywords ──
        if optimized.get('summary'):
            enhanced = self._enhance_summary(
                optimized['summary'], jd_normalized, all_jd_skills
            )
            if enhanced != optimized['summary']:
                optimized['summary'] = enhanced
                changes.append("Enhanced summary with JD keywords")

        # ── 5) Mark optimization metadata ──
        optimized['_optimization'] = {
            'optimized_at': datetime.now().isoformat(),
            'target_job': job.get('title', 'Unknown'),
            'target_company': job.get('company', 'Unknown'),
            'changes': changes,
            'jd_keywords_count': len(jd_normalized),
        }

        # Store target for builder naming
        optimized['_target_company'] = job.get('company', 'General')
        optimized['_target_title'] = job.get('title', 'Resume')

        logger.info(f"Resume optimized — {len(changes)} changes applied")
        for c in changes:
            logger.debug(f"  → {c}")

        return optimized

    # ════════════════════════════════════════════════════════════
    #  PUBLIC: check_formatting()
    # ════════════════════════════════════════════════════════════

    def check_formatting(self, docx_path: str) -> dict:
        """
        Check DOCX file for ATS formatting compliance.
        
        Analyzes:
        - Font consistency (ATS-safe fonts)
        - Section header detection
        - Bullet point format
        - Page count / document length
        - Table usage (ATS problematic)
        - Image usage (ATS can't parse)
        - Header/footer content (some ATS skip these)
        
        Args:
            docx_path: Path to .docx file
            
        Returns:
            dict: {
                score: 0-100,
                page_estimate: int,
                word_count: int,
                fonts_used: [],
                unsafe_fonts: [],
                has_tables: bool,
                has_images: bool,
                has_headers_footers: bool,
                sections_found: [],
                issues: [],
                suggestions: [],
                details: {}
            }
        """
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed")
            return {
                'score': 0,
                'issues': ['python-docx not installed — cannot analyze DOCX'],
                'suggestions': ['pip install python-docx'],
            }

        if not os.path.exists(docx_path):
            logger.error(f"File not found: {docx_path}")
            return {
                'score': 0,
                'issues': [f'File not found: {docx_path}'],
                'suggestions': [],
            }

        try:
            doc = Document(docx_path)
        except Exception as e:
            logger.error(f"Cannot open DOCX: {e}")
            return {
                'score': 0,
                'issues': [f'Cannot open DOCX: {e}'],
                'suggestions': [],
            }

        issues = []
        suggestions = []
        score_deductions = 0  # start at 100, deduct

        # ── Collect all text ──
        all_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text.strip())

        full_text = ' '.join(all_text)
        word_count = len(full_text.split())

        # ── 1) Font analysis ──
        fonts_used = set()
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font and run.font.name:
                    fonts_used.add(run.font.name)

        # Also check default style font
        default_font = None
        try:
            default_font = doc.styles['Normal'].font.name
            if default_font:
                fonts_used.add(default_font)
        except Exception:
            pass

        fonts_used = sorted(fonts_used)
        unsafe_fonts = [
            f for f in fonts_used
            if f.lower() not in self.ATS_SAFE_FONTS
        ]

        if unsafe_fonts:
            score_deductions += 10
            issues.append(f"Non-ATS-safe fonts detected: {', '.join(unsafe_fonts)}")
            suggestions.append(f"Use ATS-safe fonts: Calibri, Arial, Times New Roman")

        if len(fonts_used) > 2:
            score_deductions += 5
            issues.append(f"Too many fonts ({len(fonts_used)}): {', '.join(fonts_used)}")
            suggestions.append("Stick to 1-2 fonts maximum for ATS consistency")

        # ── 2) Table detection ──
        has_tables = len(doc.tables) > 0
        table_count = len(doc.tables)

        if has_tables:
            score_deductions += 15
            issues.append(f"Document contains {table_count} table(s) — many ATS parsers break on tables")
            suggestions.append("Avoid tables for layout — use tabs or plain text alignment")

        # ── 3) Image detection ──
        has_images = False
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                has_images = True
                image_count += 1

        if has_images:
            score_deductions += 10
            issues.append(f"Document contains {image_count} image(s) — ATS cannot parse images")
            suggestions.append("Remove images, icons, and logos — use text only")

        # ── 4) Headers/Footers ──
        has_headers_footers = False
        header_footer_text = []
        for section in doc.sections:
            if section.header:
                ht = ' '.join(p.text for p in section.header.paragraphs).strip()
                if ht:
                    has_headers_footers = True
                    header_footer_text.append(f"Header: {ht[:50]}")
            if section.footer:
                ft = ' '.join(p.text for p in section.footer.paragraphs).strip()
                if ft:
                    has_headers_footers = True
                    header_footer_text.append(f"Footer: {ft[:50]}")

        if has_headers_footers:
            score_deductions += 5
            issues.append("Content in headers/footers — some ATS skip these entirely")
            suggestions.append("Move all important info (name, contact) into document body")

        # ── 5) Section headers ──
        sections_found = []
        for para in doc.paragraphs:
            text_lower = para.text.strip().lower()
            for section_key, aliases in ATS_SECTION_ALIASES.items():
                if text_lower in aliases:
                    sections_found.append(section_key)
                    break

        expected_sections = {'summary', 'skills', 'experience', 'education'}
        missing_sections = expected_sections - set(sections_found)

        if missing_sections:
            score_deductions += len(missing_sections) * 5
            issues.append(f"Missing standard sections: {', '.join(missing_sections)}")
            suggestions.append("Include clearly labeled sections: Summary, Skills, Experience, Education")

        # ── 6) Word count / length ──
        page_estimate = max(1, math.ceil(word_count / 450))  # ~450 words per page

        if word_count < 200:
            score_deductions += 15
            issues.append(f"Resume too short ({word_count} words)")
            suggestions.append("Aim for 400-800 words with detailed bullet points")
        elif word_count < 350:
            score_deductions += 5
            issues.append(f"Resume may be too brief ({word_count} words)")
            suggestions.append("Add more detail to experience bullets")
        elif word_count > 1200:
            score_deductions += 5
            issues.append(f"Resume may be too long ({word_count} words, ~{page_estimate} pages)")
            suggestions.append("Trim to 1-2 pages — focus on most relevant experience")

        if page_estimate > self.MAX_PAGES:
            score_deductions += 10
            issues.append(f"Estimated {page_estimate} pages — exceeds {self.MAX_PAGES} page limit")

        # ── 7) Bullet point analysis ──
        bullet_count = 0
        for para in doc.paragraphs:
            style_name = para.style.name.lower() if para.style else ''
            text = para.text.strip()
            if 'list' in style_name or 'bullet' in style_name:
                bullet_count += 1
            elif text and (text[0] in '•●○◦-–—►▸▪' or text.startswith('* ')):
                bullet_count += 1

        if bullet_count < 5:
            score_deductions += 5
            issues.append(f"Very few bullet points ({bullet_count}) — ATS and humans prefer bullets")
            suggestions.append("Use bullet points for experience and project descriptions")

        # ── 8) Font size analysis ──
        font_sizes = set()
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font and run.font.size:
                    pt_size = run.font.size.pt
                    font_sizes.add(pt_size)

        small_fonts = [s for s in font_sizes if s < 9]
        if small_fonts:
            score_deductions += 5
            issues.append(f"Font size below 9pt detected — hard to read and some ATS struggle")
            suggestions.append("Use minimum 9pt font, ideally 10-11pt for body text")

        large_fonts = [s for s in font_sizes if s > 20]
        if large_fonts:
            score_deductions += 3
            issues.append(f"Font size above 20pt detected — wastes space")

        # ── 9) Hyperlinks ──
        hyperlink_count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                # Check if run's parent has hyperlink
                parent = run._element.getparent()
                if parent is not None and parent.tag.endswith('}hyperlink'):
                    hyperlink_count += 1

        # Also check via relationships
        for rel in doc.part.rels.values():
            if 'hyperlink' in rel.reltype:
                hyperlink_count += 1

        # hyperlinks are fine for ATS, just note them
        # No deduction

        # ── Calculate final score ──
        formatting_score = max(0, min(100, 100 - score_deductions))

        # ── Rating ──
        if formatting_score >= 85:
            rating = "Excellent"
        elif formatting_score >= 70:
            rating = "Good"
        elif formatting_score >= 55:
            rating = "Fair"
        elif formatting_score >= 40:
            rating = "Poor"
        else:
            rating = "Critical"

        result = {
            'score': formatting_score,
            'rating': rating,
            'page_estimate': page_estimate,
            'word_count': word_count,
            'fonts_used': fonts_used,
            'unsafe_fonts': unsafe_fonts,
            'has_tables': has_tables,
            'table_count': table_count,
            'has_images': has_images,
            'image_count': image_count,
            'has_headers_footers': has_headers_footers,
            'header_footer_text': header_footer_text,
            'sections_found': sections_found,
            'missing_sections': list(missing_sections),
            'bullet_count': bullet_count,
            'font_sizes': sorted(font_sizes) if font_sizes else [],
            'hyperlink_count': hyperlink_count,
            'issues': issues,
            'suggestions': suggestions,
            'details': {
                'file': docx_path,
                'file_size_kb': round(os.path.getsize(docx_path) / 1024, 1),
                'paragraph_count': len(doc.paragraphs),
                'checked_at': datetime.now().isoformat(),
            },
        }

        logger.info(
            f"DOCX format check: {formatting_score}/100 ({rating}) — "
            f"{len(issues)} issues, {word_count} words, ~{page_estimate} pages"
        )
        return result

    # ════════════════════════════════════════════════════════════
    #  PUBLIC: extract_keywords()
    # ════════════════════════════════════════════════════════════

    def extract_keywords(self, text: str) -> dict:
        """
        Extract technical keywords from text (JD or resume).
        
        Uses the KNOWN_SKILLS dictionary for recognition.
        Categorizes into required_skills, preferred_skills, tools, buzzwords.
        
        Args:
            text: Job description or resume text
            
        Returns:
            dict: {
                required_skills: [],
                preferred_skills: [],
                tools: [],
                buzzwords: [],
                all_skills: [],
                raw_counts: {skill: count}
            }
        """
        if not text:
            return {
                'required_skills': [], 'preferred_skills': [],
                'tools': [], 'buzzwords': [], 'all_skills': [],
                'raw_counts': {},
            }

        text_lower = text.lower()
        found_skills = {}  # normalized → (display_name, count, context)

        # ── Pass 1: Multi-word skills (exact phrase match) ──
        for skill_key in MULTI_WORD_SKILLS:
            # Use word boundary matching
            pattern = re.escape(skill_key)
            matches = re.findall(r'\b' + pattern + r'\b', text_lower)
            if matches:
                display = KNOWN_SKILLS[skill_key]
                norm = self._normalize_skill(display)
                if norm not in found_skills:
                    found_skills[norm] = (display, len(matches), self._get_context(skill_key, text_lower))

        # ── Pass 2: Single-word skills ──
        # Tokenize carefully
        tokens = re.findall(r'[a-zA-Z0-9#+.]+', text_lower)
        token_counts = Counter(tokens)

        for token in set(tokens):
            if token in KNOWN_SKILLS:
                display = KNOWN_SKILLS[token]
                norm = self._normalize_skill(display)
                if norm not in found_skills:
                    count = token_counts[token]
                    found_skills[norm] = (display, count, self._get_context(token, text_lower))

        # ── Pass 3: Special patterns ──
        # Version patterns: Node.js 18, Python 3.x, Java 17
        version_patterns = [
            r'\b(node\.?js)\s*\d*', r'\b(react\.?js)\s*\d*',
            r'\b(vue\.?js)\s*\d*', r'\b(next\.?js)\s*\d*',
            r'\b(nuxt\.?js)\s*\d*', r'\b(express\.?js)\s*\d*',
        ]
        for pat in version_patterns:
            m = re.search(pat, text_lower)
            if m:
                base = m.group(1).replace('.', '').replace(' ', '')
                if base in KNOWN_SKILLS:
                    display = KNOWN_SKILLS[base]
                    norm = self._normalize_skill(display)
                    if norm not in found_skills:
                        found_skills[norm] = (display, 1, '')

        # ── Categorize based on context ──
        required_skills = []
        preferred_skills = []
        tools = []
        buzzwords = []

        required_indicators = [
            'required', 'must have', 'must-have', 'mandatory',
            'essential', 'minimum', 'strong knowledge', 'proficiency in',
            'experience with', 'hands-on', 'expertise in',
        ]
        preferred_indicators = [
            'preferred', 'nice to have', 'nice-to-have', 'bonus',
            'desirable', 'plus', 'good to have', 'advantageous',
            'exposure to', 'familiarity with',
        ]
        tool_categories = {
            'git', 'github', 'gitlab', 'docker', 'kubernetes',
            'jenkins', 'jira', 'postman', 'swagger', 'figma',
            'vscode', 'vs code', 'npm', 'yarn', 'webpack',
            'nginx', 'apache', 'pm2', 'linux', 'ubuntu',
        }
        buzzword_categories = {
            'agile', 'scrum', 'kanban', 'tdd', 'bdd',
            'oop', 'solid principles', 'design patterns',
            'ci/cd', 'devops',
        }

        for norm, (display, count, context) in found_skills.items():
            norm_lower = norm.lower()

            # Check if it's in a required context
            is_required = any(ind in context for ind in required_indicators)
            is_preferred = any(ind in context for ind in preferred_indicators)

            if norm_lower in tool_categories or self._normalize_skill(display).lower() in tool_categories:
                tools.append(display)
            elif norm_lower in buzzword_categories:
                buzzwords.append(display)
            elif is_required:
                required_skills.append(display)
            elif is_preferred:
                preferred_skills.append(display)
            else:
                # Default: high count → required, low → preferred
                if count >= 2:
                    required_skills.append(display)
                else:
                    preferred_skills.append(display)

        all_skills = sorted(set(
            required_skills + preferred_skills + tools + buzzwords
        ))

        raw_counts = {
            display: count
            for norm, (display, count, ctx) in found_skills.items()
        }

        return {
            'required_skills': sorted(set(required_skills)),
            'preferred_skills': sorted(set(preferred_skills)),
            'tools': sorted(set(tools)),
            'buzzwords': sorted(set(buzzwords)),
            'all_skills': all_skills,
            'raw_counts': raw_counts,
        }

    # ════════════════════════════════════════════════════════════
    #  PUBLIC: keyword_gap()
    # ════════════════════════════════════════════════════════════

    def keyword_gap(self, resume_data, job: dict) -> dict:
        """
        Detailed keyword gap analysis between resume and job.
        
        Returns:
            dict: {
                matched: [],
                missing_required: [],
                missing_preferred: [],
                missing_tools: [],
                match_percentage: float,
                gap_severity: 'low'/'medium'/'high'/'critical',
                action_items: [],
            }
        """
        data = self._to_dict(resume_data)
        jd_text = self._get_jd_text(job)

        if not jd_text:
            return {
                'matched': [], 'missing_required': [], 'missing_preferred': [],
                'missing_tools': [], 'match_percentage': 0.0,
                'gap_severity': 'critical', 'action_items': ['No JD provided'],
            }

        jd_kw = self.extract_keywords(jd_text)
        resume_text = self._resume_to_text(data)
        resume_kw = self.extract_keywords(resume_text)

        # Add explicit skills
        explicit = self._get_explicit_skills(data)
        resume_all = {self._normalize_skill(s) for s in resume_kw['all_skills']}
        resume_all.update(self._normalize_skill(s) for s in explicit)

        # Compare each category
        def find_gaps(jd_list, resume_norm_set):
            matched = []
            missing = []
            for skill in jd_list:
                if self._normalize_skill(skill) in resume_norm_set:
                    matched.append(skill)
                else:
                    missing.append(skill)
            return matched, missing

        req_matched, req_missing = find_gaps(jd_kw['required_skills'], resume_all)
        pref_matched, pref_missing = find_gaps(jd_kw['preferred_skills'], resume_all)
        tool_matched, tool_missing = find_gaps(jd_kw['tools'], resume_all)
        buzz_matched, buzz_missing = find_gaps(jd_kw['buzzwords'], resume_all)

        all_matched = sorted(set(req_matched + pref_matched + tool_matched + buzz_matched))
        all_jd = sorted(set(
            jd_kw['required_skills'] + jd_kw['preferred_skills'] +
            jd_kw['tools'] + jd_kw['buzzwords']
        ))

        match_pct = (len(all_matched) / len(all_jd) * 100) if all_jd else 0

        if match_pct >= 80:
            severity = 'low'
        elif match_pct >= 60:
            severity = 'medium'
        elif match_pct >= 40:
            severity = 'high'
        else:
            severity = 'critical'

        # Build action items
        action_items = []
        if req_missing:
            action_items.append(
                f"🔴 Add required skills to resume: {', '.join(req_missing[:6])}"
            )
        if pref_missing:
            action_items.append(
                f"🟡 Consider adding preferred skills: {', '.join(pref_missing[:5])}"
            )
        if tool_missing:
            action_items.append(
                f"🔧 Add tools if you've used them: {', '.join(tool_missing[:5])}"
            )
        if not action_items:
            action_items.append("✅ Great keyword coverage — no critical gaps")

        return {
            'matched': all_matched,
            'missing_required': req_missing,
            'missing_preferred': pref_missing,
            'missing_tools': tool_missing,
            'missing_buzzwords': buzz_missing,
            'match_percentage': round(match_pct, 1),
            'gap_severity': severity,
            'action_items': action_items,
            'jd_breakdown': {
                'required': len(jd_kw['required_skills']),
                'preferred': len(jd_kw['preferred_skills']),
                'tools': len(jd_kw['tools']),
                'buzzwords': len(jd_kw['buzzwords']),
                'total': len(all_jd),
            },
        }

    # ════════════════════════════════════════════════════════════
    #  INTERNAL: Scoring sub-functions
    # ════════════════════════════════════════════════════════════

    def _calc_density_score(self, section_text: str, jd_normalized: Set[str]) -> float:
        """Score keyword density in a section vs JD keywords."""
        if not section_text or not jd_normalized:
            return 0.0

        section_lower = section_text.lower()
        found = 0
        for skill_norm in jd_normalized:
            # Check if normalized skill appears in section
            skill_lower = skill_norm.lower()
            if skill_lower in section_lower:
                found += 1
            else:
                # Try common variations
                variations = [
                    skill_lower.replace('.', ''),
                    skill_lower.replace('.js', ''),
                    skill_lower.replace(' ', ''),
                ]
                if any(v in section_lower for v in variations if v):
                    found += 1

        if jd_normalized:
            return min(100, (found / len(jd_normalized)) * 100)
        return 0.0

    def _calc_section_order_score(self, data: dict) -> float:
        """Score section ordering vs recommended ATS order."""
        present_sections = []
        if data.get('summary'):
            present_sections.append('summary')
        if data.get('skills'):
            present_sections.append('skills')
        if data.get('experience'):
            present_sections.append('experience')
        if data.get('projects'):
            present_sections.append('projects')
        if data.get('education'):
            present_sections.append('education')
        if data.get('certifications'):
            present_sections.append('certifications')
        if data.get('achievements'):
            present_sections.append('achievements')

        if len(present_sections) < 2:
            return 100.0

        # Check if order matches recommended
        ideal = [s for s in ATS_SECTION_ORDER if s in present_sections]
        if present_sections == ideal:
            return 100.0

        # Measure disorder using inversions
        inversions = 0
        total_pairs = 0
        for i in range(len(present_sections)):
            for j in range(i + 1, len(present_sections)):
                total_pairs += 1
                s1 = present_sections[i]
                s2 = present_sections[j]
                if s1 in ATS_SECTION_ORDER and s2 in ATS_SECTION_ORDER:
                    idx1 = ATS_SECTION_ORDER.index(s1)
                    idx2 = ATS_SECTION_ORDER.index(s2)
                    if idx1 > idx2:
                        inversions += 1

        if total_pairs == 0:
            return 100.0

        return max(0, 100 - (inversions / total_pairs) * 100)

    def _calc_text_keyword_score(self, text: str, jd_normalized: Set[str]) -> float:
        """Score how many JD keywords appear in a text block."""
        if not text or not jd_normalized:
            return 0.0

        text_lower = text.lower()
        found = 0

        for skill_norm in jd_normalized:
            skill_lower = skill_norm.lower()
            if skill_lower in text_lower:
                found += 1
            else:
                # Try variations
                variations = [
                    skill_lower.replace('.', ''),
                    skill_lower.replace('.js', ''),
                    skill_lower.replace(' ', ''),
                    skill_lower.replace('-', ''),
                ]
                if any(v in text_lower for v in variations if len(v) >= 2):
                    found += 1

        return min(100, (found / len(jd_normalized)) * 100)

    def _calc_formatting_score(self, data: dict) -> Tuple[float, List[str], List[str]]:
        """Score resume data structure for ATS-friendliness."""
        issues = []
        suggestions = []
        score_deductions = 0

        # Check summary length
        summary = data.get('summary', '')
        if not summary:
            score_deductions += 15
            issues.append("No professional summary section")
            suggestions.append("Add a 2-3 sentence professional summary")
        elif len(summary.split()) < 15:
            score_deductions += 5
            issues.append("Summary too short")
            suggestions.append("Expand summary to 30-60 words with key skills")
        elif len(summary.split()) > 80:
            score_deductions += 3
            issues.append("Summary too long")
            suggestions.append("Trim summary to 40-60 words")

        # Check skills section
        skills = data.get('skills', {})
        if not skills:
            score_deductions += 15
            issues.append("No skills section")
            suggestions.append("Add a categorized skills section")
        elif isinstance(skills, dict):
            total_skills = sum(
                len(v) if isinstance(v, list) else 1
                for v in skills.values()
            )
            if total_skills < 5:
                score_deductions += 5
                issues.append(f"Very few skills listed ({total_skills})")
                suggestions.append("List at least 10-15 relevant technical skills")

        # Check experience
        experience = data.get('experience', [])
        if not experience:
            score_deductions += 15
            issues.append("No experience section")
        else:
            for i, exp in enumerate(experience):
                bullets = self._extract_bullets_from_entry(exp)
                if len(bullets) < self.MIN_BULLETS_PER_EXP:
                    score_deductions += 3
                    issues.append(
                        f"Experience '{exp.get('title', f'#{i+1}')}' has "
                        f"only {len(bullets)} bullet(s) — minimum {self.MIN_BULLETS_PER_EXP}"
                    )
                elif len(bullets) > self.MAX_BULLETS_PER_EXP:
                    score_deductions += 2
                    issues.append(
                        f"Experience '{exp.get('title', f'#{i+1}')}' has "
                        f"{len(bullets)} bullets — trim to {self.MAX_BULLETS_PER_EXP}"
                    )

                # Check for action verbs
                weak_starts = 0
                for b in bullets:
                    first_word = b.split()[0].lower() if b.split() else ''
                    if first_word in ('i', 'the', 'a', 'an', 'my', 'we', 'our', 'this', 'was', 'is'):
                        weak_starts += 1
                if weak_starts > 0:
                    score_deductions += 2
                    suggestions.append(
                        f"Start bullets with action verbs (Built, Developed, Designed, Implemented)"
                    )

        # Check education
        if not data.get('education'):
            score_deductions += 5
            issues.append("No education section")

        # Overall word count estimate
        full_text = self._resume_to_text(data)
        word_count = len(full_text.split())
        if word_count < self.IDEAL_WORD_COUNT[0]:
            score_deductions += 5
            issues.append(f"Resume too brief ({word_count} words)")
            suggestions.append(f"Aim for {self.IDEAL_WORD_COUNT[0]}-{self.IDEAL_WORD_COUNT[1]} words")
        elif word_count > self.IDEAL_WORD_COUNT[1]:
            score_deductions += 3
            suggestions.append("Consider trimming less relevant details")

        formatting_score = max(0, min(100, 100 - score_deductions))
        return formatting_score, issues, suggestions

    # ════════════════════════════════════════════════════════════
    #  INTERNAL: Optimization helpers
    # ════════════════════════════════════════════════════════════

    def _reorder_skills(self, skills, jd_normalized: Set[str]) -> Union[dict, list]:
        """Reorder skills categories — put matching skills/categories first."""
        if not skills:
            return skills

        if isinstance(skills, list):
            # Simple list: matching first
            matching = []
            non_matching = []
            for s in skills:
                if self._normalize_skill(str(s)) in jd_normalized:
                    matching.append(s)
                else:
                    non_matching.append(s)
            return matching + non_matching

        if isinstance(skills, dict):
            # Dict: score each category by how many JD matches it has
            import copy
            scored_cats = []
            for cat, items in skills.items():
                if isinstance(items, list):
                    match_count = sum(
                        1 for item in items
                        if self._normalize_skill(str(item)) in jd_normalized
                    )
                    # Reorder items within category too
                    matching = [
                        i for i in items
                        if self._normalize_skill(str(i)) in jd_normalized
                    ]
                    non_matching = [
                        i for i in items
                        if self._normalize_skill(str(i)) not in jd_normalized
                    ]
                    scored_cats.append((cat, matching + non_matching, match_count))
                else:
                    scored_cats.append((cat, items, 0))

            # Sort categories: most matches first
            scored_cats.sort(key=lambda x: x[2], reverse=True)

            from collections import OrderedDict
            result = OrderedDict()
            for cat, items, _ in scored_cats:
                result[cat] = items
            return dict(result)

        return skills

    def _reorder_bullets(self, experience: dict, jd_normalized: Set[str]) -> dict:
        """Reorder experience bullets — JD-relevant bullets first."""
        import copy
        exp = copy.deepcopy(experience)

        desc = exp.get('description', '')
        if not desc:
            return exp

        if isinstance(desc, list):
            bullets = [str(b) for b in desc]
        elif isinstance(desc, str):
            bullets = [b.strip() for b in desc.split('\n') if b.strip()]
        else:
            return exp

        if not bullets:
            return exp

        # Score each bullet by JD keyword presence
        scored = []
        for bullet in bullets:
            bullet_lower = bullet.lower()
            score = sum(
                1 for kw in jd_normalized
                if kw.lower() in bullet_lower
                or kw.lower().replace('.', '') in bullet_lower
                or kw.lower().replace('.js', '') in bullet_lower
            )
            scored.append((bullet, score))

        # Stable sort: matching first, preserve relative order otherwise
        scored.sort(key=lambda x: x[1], reverse=True)
        exp['description'] = [b for b, s in scored]

        return exp

    def _reorder_projects(self, projects: list, jd_normalized: Set[str]) -> list:
        """Reorder projects — most JD-relevant first."""
        if not projects:
            return projects

        import copy

        scored = []
        for proj in projects:
            proj_text = ''
            tech = proj.get('tech', proj.get('technologies', proj.get('stack', '')))
            if isinstance(tech, list):
                proj_text = ' '.join(str(t) for t in tech)
            elif tech:
                proj_text = str(tech)

            desc = proj.get('description', '')
            if isinstance(desc, list):
                proj_text += ' ' + ' '.join(str(d) for d in desc)
            elif desc:
                proj_text += ' ' + str(desc)

            proj_lower = proj_text.lower()
            match_count = sum(
                1 for kw in jd_normalized
                if kw.lower() in proj_lower
                or kw.lower().replace('.', '') in proj_lower
            )
            scored.append((copy.deepcopy(proj), match_count))

        scored.sort(key=lambda x: x[1], reverse=True)
        return [p for p, s in scored]

    def _enhance_summary(
        self,
        summary: str,
        jd_normalized: Set[str],
        jd_display_skills: Set[str],
    ) -> str:
        """
        Enhance summary to include more JD keywords.
        Only adds skills that are genuinely in the resume elsewhere.
        NEVER fabricates.
        """
        if not summary or not jd_normalized:
            return summary

        summary_lower = summary.lower()

        # Find JD skills NOT yet in summary
        missing_in_summary = []
        for skill_display in jd_display_skills:
            norm = self._normalize_skill(skill_display)
            skill_lower = norm.lower()
            if skill_lower not in summary_lower and \
               skill_lower.replace('.', '') not in summary_lower:
                missing_in_summary.append(skill_display)

        if not missing_in_summary or len(missing_in_summary) > 10:
            # Don't over-stuff
            return summary

        # Only add top 3-4 most important missing keywords
        # (prioritize shorter skill names — they're usually core techs)
        missing_in_summary.sort(key=len)
        to_add = missing_in_summary[:4]

        if not to_add:
            return summary

        # Check if summary ends with period
        summary = summary.rstrip()
        if not summary.endswith('.'):
            summary += '.'

        # Append a natural clause
        skills_text = ', '.join(to_add)
        suffix = f" Experienced with {skills_text}."
        enhanced = summary + suffix

        # Don't make it too long
        if len(enhanced.split()) > 80:
            return summary

        return enhanced

    # ════════════════════════════════════════════════════════════
    #  INTERNAL: Utility helpers
    # ════════════════════════════════════════════════════════════

    def _to_dict(self, resume_data) -> dict:
        """Convert ResumeData to dict."""
        if isinstance(resume_data, dict):
            return resume_data
        try:
            from dataclasses import asdict
            return asdict(resume_data)
        except (TypeError, ImportError):
            pass
        # Fallback: attribute extraction
        fields = [
            'name', 'email', 'phone', 'location', 'linkedin', 'github',
            'summary', 'education', 'experience', 'projects', 'skills',
            'certifications', 'achievements', 'coding_profiles',
        ]
        result = {}
        for f in fields:
            val = getattr(resume_data, f, None)
            if val is not None:
                result[f] = val
        return result

    def _get_jd_text(self, job: dict) -> str:
        """Extract all text from job dict for keyword extraction."""
        parts = []
        if job.get('description'):
            parts.append(str(job['description']))
        if job.get('title'):
            parts.append(str(job['title']))
        if job.get('skills'):
            if isinstance(job['skills'], list):
                parts.append(' '.join(str(s) for s in job['skills']))
            else:
                parts.append(str(job['skills']))
        return ' '.join(parts)

    def _resume_to_text(self, data: dict) -> str:
        """Flatten entire resume data to plain text."""
        parts = []

        if data.get('summary'):
            parts.append(str(data['summary']))

        # Skills
        skills = data.get('skills', {})
        if isinstance(skills, dict):
            for cat, items in skills.items():
                parts.append(str(cat))
                if isinstance(items, list):
                    parts.extend(str(i) for i in items)
                else:
                    parts.append(str(items))
        elif isinstance(skills, list):
            parts.extend(str(s) for s in skills)

        # Experience
        for exp in data.get('experience', []):
            parts.append(str(exp.get('title', '')))
            parts.append(str(exp.get('company', '')))
            if exp.get('stack'):
                stack = exp['stack']
                if isinstance(stack, list):
                    parts.extend(str(s) for s in stack)
                else:
                    parts.append(str(stack))
            desc = exp.get('description', '')
            if isinstance(desc, list):
                parts.extend(str(d) for d in desc)
            elif desc:
                parts.append(str(desc))
            if exp.get('projects'):
                if isinstance(exp['projects'], list):
                    for p in exp['projects']:
                        parts.append(str(p) if not isinstance(p, dict) else str(p.get('name', '')))
                else:
                    parts.append(str(exp['projects']))

        # Projects
        for proj in data.get('projects', []):
            parts.append(str(proj.get('name', '')))
            tech = proj.get('tech', proj.get('technologies', proj.get('stack', '')))
            if isinstance(tech, list):
                parts.extend(str(t) for t in tech)
            elif tech:
                parts.append(str(tech))
            desc = proj.get('description', '')
            if isinstance(desc, list):
                parts.extend(str(d) for d in desc)
            elif desc:
                parts.append(str(desc))

        # Education
        for edu in data.get('education', []):
            parts.append(str(edu.get('degree', '')))
            parts.append(str(edu.get('university', edu.get('institution', ''))))

        # Certifications
        for cert in data.get('certifications', []):
            if isinstance(cert, dict):
                parts.append(str(cert.get('name', '')))
            else:
                parts.append(str(cert))

        # Achievements
        for ach in data.get('achievements', []):
            parts.append(str(ach))

        return ' '.join(p for p in parts if p)

    def _get_section_text(self, data: dict, section: str) -> str:
        """Get flattened text for a specific resume section."""
        if section == 'skills':
            skills = data.get('skills', {})
            if isinstance(skills, dict):
                parts = []
                for cat, items in skills.items():
                    parts.append(str(cat))
                    if isinstance(items, list):
                        parts.extend(str(i) for i in items)
                    else:
                        parts.append(str(items))
                return ' '.join(parts)
            elif isinstance(skills, list):
                return ' '.join(str(s) for s in skills)
            return ''

        elif section == 'experience':
            parts = []
            for exp in data.get('experience', []):
                parts.append(str(exp.get('title', '')))
                parts.append(str(exp.get('company', '')))
                desc = exp.get('description', '')
                if isinstance(desc, list):
                    parts.extend(str(d) for d in desc)
                elif desc:
                    parts.append(str(desc))
                if exp.get('stack'):
                    stack = exp['stack']
                    if isinstance(stack, list):
                        parts.extend(str(s) for s in stack)
                    else:
                        parts.append(str(stack))
            return ' '.join(parts)

        elif section == 'summary':
            return str(data.get('summary', ''))

        return ''

    def _get_explicit_skills(self, data: dict) -> Set[str]:
        """Get all skills explicitly listed in the skills section."""
        skills = data.get('skills', {})
        result = set()

        if isinstance(skills, dict):
            for cat, items in skills.items():
                if isinstance(items, list):
                    result.update(str(i) for i in items)
                elif items:
                    result.add(str(items))
        elif isinstance(skills, list):
            result.update(str(s) for s in skills)

        return result

    @staticmethod
    def _normalize_skill(skill: str) -> str:
        """Normalize skill name for comparison."""
        if not skill:
            return ''
        s = str(skill).strip().lower()
        # Remove version numbers
        s = re.sub(r'\s*\d+(\.\d+)*\s*$', '', s)
        # Remove trailing parens content: "JavaScript (ES6+)" → "javascript"
        s = re.sub(r'\s*\([^)]*\)\s*$', '', s)
        # Normalize .js variations
        s = s.replace('.js', 'js').replace('node js', 'nodejs')
        # Remove extra spaces
        s = re.sub(r'\s+', ' ', s).strip()
        return s

    @staticmethod
    def _get_context(keyword: str, text: str, window: int = 100) -> str:
        """Get surrounding context for a keyword in text."""
        idx = text.find(keyword)
        if idx == -1:
            return ''
        start = max(0, idx - window)
        end = min(len(text), idx + len(keyword) + window)
        return text[start:end].lower()

    @staticmethod
    def _extract_bullets_from_entry(entry: dict) -> list:
        """Extract bullet point texts from an experience/project entry."""
        desc = entry.get('description', '')
        if not desc:
            return []
        if isinstance(desc, list):
            return [str(b).strip() for b in desc if b]
        if isinstance(desc, str):
            return [b.strip() for b in desc.split('\n') if b.strip()]
        return []

    def _empty_score(self, reason: str) -> dict:
        """Return empty score dict with reason."""
        return {
            'score': 0, 'grade': 'F',
            'keyword_score': 0, 'density_score': 0,
            'section_score': 0, 'summary_score': 0,
            'experience_score': 0, 'formatting_score': 0,
            'matched_keywords': [], 'missing_keywords': [],
            'total_jd_keywords': 0, 'match_ratio': '0/0',
            'issues': [reason], 'suggestions': [],
            'details': {},
        }


# ════════════════════════════════════════════════════════════════
#  TEST BLOCK
# ════════════════════════════════════════════════════════════════

if __name__ == "__main__":

    print("=" * 70)
    print("  ATS RESUME OPTIMIZER — Test Suite")
    print("=" * 70)

    # ── Load resume data ──
    try:
        from profile.resume_data import get_base_resume, resume_to_dict
        base_resume = get_base_resume()
        resume_dict = resume_to_dict(base_resume)
        print("\n✅ Loaded resume from profile/resume_data.py")
    except Exception as e:
        print(f"\n⚠️  Could not load profile/resume_data.py ({e})")
        print("   Using built-in test data.\n")
        resume_dict = {
            'name': 'Piyush Kashyap',
            'email': 'piyushkashyap3247@gmail.com',
            'phone': '+91 73107 03247',
            'location': 'Rishikesh, Uttarakhand, India',
            'linkedin': 'linkedin.com/in/piyush-kashyap731',
            'github': 'github.com/Piyush731',
            'summary': (
                'Full Stack Developer L1 with end-to-end responsibility of 10+ production '
                'applications across fintech, ERP, edtech, and CRM domains. Sole developer '
                'on each project. Shipped multi-tenant systems with 57+ DB tables, real-time '
                'WebSocket feeds, and third-party integrations.'
            ),
            'skills': {
                'Languages': ['JavaScript (ES6+)', 'Java', 'Python', 'SQL'],
                'Frontend': ['Vue.js', 'Nuxt.js', 'React.js', 'Vuetify', 'Tailwind CSS'],
                'Backend': ['Node.js', 'Express.js', 'Spring Boot', 'REST APIs', 'WebSockets', 'Microservices'],
                'Databases': ['MySQL', 'MongoDB', 'PostgreSQL', 'Redis'],
                'Tools': ['Git', 'Docker', 'Kafka', 'JWT', 'Playwright'],
            },
            'experience': [
                {
                    'title': 'Full Stack Developer L1',
                    'company': 'Site Guru Pvt Ltd',
                    'duration': 'Aug 2024 — Present',
                    'stack': 'Vue.js, Nuxt.js, Node.js, MySQL, Vuetify, REST APIs, WebSockets',
                    'description': [
                        'Sole developer on 10+ production apps across fintech, ERP, edtech, CRM',
                        'Built BizHub ERP with 57 DB tables, multi-tenant architecture, 5 user types',
                        'Developed My RTO Expert handling 1000+ daily requests with META WhatsApp API',
                        'Created Rudra Fintech investment platform with interest/TDS calculation',
                        'Built FX Prime Trading with IOTrades API, MT5, WebSocket price feeds',
                        'Shipped Dheeranet ISP CRM with hierarchical zone management',
                    ],
                    'recognition': ['Top performer', 'Production deployment authority'],
                },
                {
                    'title': 'Salesforce Developer Intern',
                    'company': 'SmartBridge',
                    'duration': 'July — Sept 2024',
                    'description': [
                        'Built Apex triggers and batch classes for lead processing',
                        'Developed Lightning Web Components for custom UI',
                        'Automated lead assignment reducing manual work by 30%',
                    ],
                },
            ],
            'projects': [
                {
                    'name': 'Collaborative Workspace',
                    'tech': ['React', 'Node.js', 'MongoDB', 'Socket.io', 'Gitea API'],
                    'description': ['Real-time collaborative platform with RBAC for 50+ users'],
                },
                {
                    'name': 'Invoice Microservice',
                    'tech': ['Java', 'Spring Boot', 'PostgreSQL', 'Kafka', 'Docker'],
                    'description': ['Event-driven microservice for invoice processing with PDF export'],
                },
                {
                    'name': 'CareerCraft AI Resume Analyzer',
                    'tech': ['Python', 'Gemini API', 'Streamlit'],
                    'description': ['NLP-powered resume analysis with JD match scoring'],
                },
            ],
            'education': [
                {
                    'degree': 'B.Tech Computer Science',
                    'university': 'Graphic Era Hill University',
                    'year': '2021 — 2025',
                    'cgpa': '7.79/10',
                },
            ],
            'certifications': [
                {'name': 'Full Stack Java Developer', 'issuer': 'Udemy'},
                {'name': 'Agile Project Management', 'issuer': 'Udemy'},
            ],
            'achievements': [
                '100+ problems on LeetCode/GFG',
                'Published TutorsUp on Google Play Store',
                'Top performer — production deployment authority',
            ],
        }

    optimizer = ATSOptimizer()

    # ── Test JD: Java/Spring Boot Backend Developer ──
    test_job_java = {
        'title': 'Backend Developer - Java',
        'company': 'Razorpay',
        'location': 'Bangalore',
        'description': """
        We are looking for a Backend Developer with strong Java and Spring Boot experience.

        Required Skills:
        - 1-3 years experience in Java, Spring Boot, Microservices
        - Strong understanding of REST APIs and system design
        - Experience with PostgreSQL or MySQL databases
        - Hands-on with Docker, Kubernetes, CI/CD pipelines
        - Understanding of Kafka or RabbitMQ for event-driven architecture
        - Git version control, Agile/Scrum methodology
        - Data structures and algorithms proficiency

        Preferred:
        - Experience with Redis caching
        - Familiarity with AWS or GCP cloud services
        - Exposure to React.js or Angular frontend
        - GraphQL API development
        - Monitoring tools: Grafana, Prometheus

        Responsibilities:
        - Design and develop scalable backend services
        - Write clean, testable code with JUnit tests
        - Participate in code reviews and system design discussions
        - Optimize database queries and application performance
        - Deploy and monitor services in production
        """,
        'skills': ['Java', 'Spring Boot', 'PostgreSQL', 'Docker', 'Kafka', 'Microservices'],
    }

    # ── Test JD: MERN Full Stack Developer ──
    test_job_mern = {
        'title': 'Full Stack Developer (MERN)',
        'company': 'Zomato',
        'location': 'Gurugram',
        'description': """
        Looking for a Full Stack Developer experienced with the MERN stack.

        Requirements:
        - Proficiency in React.js, Node.js, Express.js, MongoDB
        - Experience building REST APIs and WebSocket applications
        - Frontend: HTML5, CSS3, Tailwind CSS or Bootstrap
        - State management: Redux or Context API
        - Testing: Jest, Cypress
        - Git, GitHub, Agile development
        - Strong problem solving and DSA skills

        Nice to have:
        - Next.js or Nuxt.js experience
        - TypeScript
        - Docker containerization
        - Redis caching
        - AWS deployment experience
        - GraphQL
        """,
        'skills': ['React.js', 'Node.js', 'MongoDB', 'Express.js', 'TypeScript'],
    }

    # ═══════════════════════════════════════════════════════════
    #  Test 1: Keyword Extraction
    # ═══════════════════════════════════════════════════════════
    print("\n─── Test 1: Keyword Extraction ─────────────────────────")
    kw = optimizer.extract_keywords(test_job_java['description'])
    print(f"  Required : {kw['required_skills']}")
    print(f"  Preferred: {kw['preferred_skills']}")
    print(f"  Tools    : {kw['tools']}")
    print(f"  Buzzwords: {kw['buzzwords']}")
    print(f"  Total    : {len(kw['all_skills'])} skills found")

    # ═══════════════════════════════════════════════════════════
    #  Test 2: ATS Score — Java Backend
    # ═══════════════════════════════════════════════════════════
    print("\n─── Test 2: ATS Score — Java Backend (Razorpay) ────────")
    score_java = optimizer.score(resume_dict, test_job_java)
    print(f"  Overall Score: {score_java['score']}/100 (Grade: {score_java['grade']})")
    print(f"  Keyword Match: {score_java['keyword_score']}/100 ({score_java['match_ratio']})")
    print(f"  Skill Density: {score_java['density_score']}/100")
    print(f"  Section Order: {score_java['section_score']}/100")
    print(f"  Summary Score: {score_java['summary_score']}/100")
    print(f"  Exp Keywords : {score_java['experience_score']}/100")
    print(f"  Formatting   : {score_java['formatting_score']}/100")
    print(f"  ✅ Matched   : {', '.join(score_java['matched_keywords'][:10])}")
    print(f"  ❌ Missing   : {', '.join(score_java['missing_keywords'][:10])}")
    if score_java['issues']:
        print(f"  Issues:")
        for issue in score_java['issues'][:5]:
            print(f"    ⚠ {issue}")
    if score_java['suggestions']:
        print(f"  Suggestions:")
        for sug in score_java['suggestions'][:5]:
            print(f"    💡 {sug}")

    # ═══════════════════════════════════════════════════════════
    #  Test 3: ATS Score — MERN Full Stack
    # ═══════════════════════════════════════════════════════════
    print("\n─── Test 3: ATS Score — MERN Full Stack (Zomato) ───────")
    score_mern = optimizer.score(resume_dict, test_job_mern)
    print(f"  Overall Score: {score_mern['score']}/100 (Grade: {score_mern['grade']})")
    print(f"  Keyword Match: {score_mern['keyword_score']}/100 ({score_mern['match_ratio']})")
    print(f"  ✅ Matched   : {', '.join(score_mern['matched_keywords'][:10])}")
    print(f"  ❌ Missing   : {', '.join(score_mern['missing_keywords'][:10])}")

    # ═══════════════════════════════════════════════════════════
    #  Test 4: Keyword Gap Analysis
    # ═══════════════════════════════════════════════════════════
    print("\n─── Test 4: Keyword Gap Analysis ────────────────────────")
    gap = optimizer.keyword_gap(resume_dict, test_job_java)
    print(f"  Match %      : {gap['match_percentage']}%")
    print(f"  Severity     : {gap['gap_severity']}")
    print(f"  Matched      : {len(gap['matched'])} skills")
    print(f"  Missing Req  : {gap['missing_required']}")
    print(f"  Missing Pref : {gap['missing_preferred']}")
    print(f"  Missing Tools: {gap['missing_tools']}")
    print(f"  JD Breakdown : {gap['jd_breakdown']}")
    for action in gap['action_items']:
        print(f"  → {action}")

    # ═══════════════════════════════════════════════════════════
    #  Test 5: Optimize Resume
    # ═══════════════════════════════════════════════════════════
    print("\n─── Test 5: Optimize Resume ─────────────────────────────")
    optimized = optimizer.optimize(resume_dict, test_job_java)
    opt_meta = optimized.get('_optimization', {})
    print(f"  Target: {opt_meta.get('target_job')} @ {opt_meta.get('target_company')}")
    print(f"  Changes applied:")
    for change in opt_meta.get('changes', []):
        print(f"    ✅ {change}")

    # Show skills reordering
    orig_skills = resume_dict.get('skills', {})
    opt_skills = optimized.get('skills', {})
    if isinstance(orig_skills, dict) and isinstance(opt_skills, dict):
        orig_cats = list(orig_skills.keys())
        opt_cats = list(opt_skills.keys())
        if orig_cats != opt_cats:
            print(f"  Skills categories reordered:")
            print(f"    Before: {orig_cats}")
            print(f"    After : {opt_cats}")

    # Score optimized vs original
    print("\n  Score comparison:")
    score_before = optimizer.score(resume_dict, test_job_java)
    score_after = optimizer.score(optimized, test_job_java)
    print(f"    Before optimization: {score_before['score']}/100")
    print(f"    After optimization : {score_after['score']}/100")
    diff = score_after['score'] - score_before['score']
    if diff > 0:
        print(f"    ↑ Improvement: +{diff} points")
    elif diff == 0:
        print(f"    = No change (already well-optimized)")
    else:
        print(f"    ↓ Decrease: {diff} points (check optimization logic)")

     # ═══════════════════════════════════════════════════════════
    #  Test 6: Check DOCX Formatting
    # ═══════════════════════════════════════════════════════════
    print("\n─── Test 6: Check DOCX Formatting ──────────────────────")
    # Try to find a generated DOCX from builder tests
    from pathlib import Path

    # Resolve output dir from RESUME_CONFIG or fallback
    _output_dir = RESUME_CONFIG.get('output_dir', os.path.join(BASE_DIR, 'resume', 'output'))
    docx_files = list(Path(_output_dir).rglob('*.docx'))

    if docx_files:
        docx_path = str(docx_files[0])
        print(f"  Checking: {docx_path}")
        fmt_result = optimizer.check_formatting(docx_path)
        print(f"  Format Score: {fmt_result['score']}/100 ({fmt_result.get('rating', 'N/A')})")
        print(f"  Pages (est) : {fmt_result.get('page_estimate', '?')}")
        print(f"  Word count  : {fmt_result.get('word_count', '?')}")
        print(f"  Fonts used  : {fmt_result.get('fonts_used', [])}")
        print(f"  Unsafe fonts: {fmt_result.get('unsafe_fonts', [])}")
        print(f"  Has tables  : {fmt_result.get('has_tables', False)}")
        print(f"  Has images  : {fmt_result.get('has_images', False)}")
        print(f"  Sections    : {fmt_result.get('sections_found', [])}")
        print(f"  Missing     : {fmt_result.get('missing_sections', [])}")
        print(f"  Bullets     : {fmt_result.get('bullet_count', 0)}")
        if fmt_result.get('issues'):
            print(f"  Issues:")
            for issue in fmt_result['issues'][:5]:
                print(f"    ⚠ {issue}")
        if fmt_result.get('suggestions'):
            print(f"  Suggestions:")
            for sug in fmt_result['suggestions'][:5]:
                print(f"    💡 {sug}")
    else:
        print(f"  ℹ️  No DOCX files found in {_output_dir}")
        print(f"     Run resume/builder.py first to generate test files")
        print(f"     Then re-run this test.")

    # ═══════════════════════════════════════════════════════════
    #  Test 7: Edge Cases
    # ═══════════════════════════════════════════════════════════
    print("\n─── Test 7: Edge Cases ─────────────────────────────────")

    # Empty JD
    empty_score = optimizer.score(resume_dict, {})
    print(f"  Empty JD     : {empty_score['score']}/100 — {empty_score['issues'][0]}")

    # Empty resume
    empty_resume_score = optimizer.score({}, test_job_java)
    print(f"  Empty resume : {empty_resume_score['score']}/100")

    # Keyword extraction from short text
    short_kw = optimizer.extract_keywords("Need Java developer with Spring Boot and React")
    print(f"  Short text KW: {short_kw['all_skills']}")

    # Nonexistent DOCX
    bad_fmt = optimizer.check_formatting("/nonexistent/path/resume.docx")
    print(f"  Bad DOCX path: score={bad_fmt['score']} — {bad_fmt['issues'][0]}")

    # ═══════════════════════════════════════════════════════════
    print("\n" + "=" * 70)
    print("  ✅ resume/optimizer.py — All tests complete")
    print("=" * 70)